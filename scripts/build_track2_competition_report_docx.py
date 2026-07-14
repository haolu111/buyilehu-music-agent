from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path.home() / "Desktop" / "第一版音乐游戏生成智能体-赛道二项目汇报稿.docx"

FONT_LATIN = "Arial"
FONT_CN = "PingFang SC"
INK = "1E293B"
MUTED = "64748B"
BLUE = "1D4ED8"
BLUE_DARK = "1E3A8A"
GREEN = "047857"
AMBER = "B45309"
ROSE = "BE123C"
VIOLET = "6D28D9"
FILL_BLUE = "EFF6FF"
FILL_GREEN = "ECFDF5"
FILL_AMBER = "FFFBEB"
FILL_VIOLET = "F5F3FF"
FILL_ROSE = "FFF1F2"
FILL_SLATE = "F8FAFC"
BORDER = "CBD5E1"
TABLE_HEADER = "E2E8F0"
PAGE_WIDTH_DXA = 9360


def set_run_font(run, size: float, *, bold: bool = False, color: str = INK, italic: bool = False) -> None:
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph(paragraph, *, before: float = 0, after: float = 6, line: float = 1.15) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 100, start: int = 140, bottom: int = 100, end: int = 140) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = PAGE_WIDTH_DXA, indent_dxa: int = 120) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_column_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def mark_row_as_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.82)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(10.8)
    normal.font.color.rgb = RGBColor.from_string(INK)

    for style_name, size, color in [
        ("Heading 1", 16, BLUE_DARK),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, INK),
    ]:
        style = styles[style_name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(14 if style_name == "Heading 1" else 9)
        style.paragraph_format.space_after = Pt(7 if style_name == "Heading 1" else 5)
        style.paragraph_format.line_spacing = 1.15

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph(header, after=0)
    run = header.add_run("赛道二：大语言模型创新应用开发 | 不亦乐乎音乐游戏生成智能体")
    set_run_font(run, 9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(footer, after=0)
    run = footer.add_run("技术是手段，育人是目的")
    set_run_font(run, 9, color=MUTED)


def heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def para(doc: Document, text: str, *, bold: bool = False, color: str = INK, after: float = 6) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, after=after)
    run = p.add_run(text)
    set_run_font(run, 10.8, bold=bold, color=color)


def bullet(doc: Document, text: str, *, prefix: str = "") -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    set_paragraph(p, after=4, line=1.15)
    if prefix and text.startswith(prefix):
        run = p.add_run(prefix)
        set_run_font(run, 10.6, bold=True)
        run = p.add_run(text[len(prefix) :])
        set_run_font(run, 10.6)
    else:
        run = p.add_run(text)
        set_run_font(run, 10.6)


def number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    set_paragraph(p, after=4, line=1.15)
    run = p.add_run(text)
    set_run_font(run, 10.6)


def callout(doc: Document, title: str, body: str, *, fill: str = FILL_BLUE, accent: str = BLUE_DARK) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table)
    mark_row_as_header(table.rows[0])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_border(cell)
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    p = cell.paragraphs[0]
    set_paragraph(p, after=3)
    run = p.add_run(title)
    set_run_font(run, 11.6, bold=True, color=accent)
    p = cell.add_paragraph()
    set_paragraph(p, after=0, line=1.15)
    run = p.add_run(body)
    set_run_font(run, 10.6)


def table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    tbl = doc.add_table(rows=1, cols=len(headers))
    set_table_width(tbl)
    mark_row_as_header(tbl.rows[0])
    for i, header in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_column_width(cell, widths[i])
        shade_cell(cell, TABLE_HEADER)
        set_cell_border(cell)
        set_cell_margins(cell, top=90, start=120, bottom=90, end=120)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        set_paragraph(p, after=0)
        run = p.add_run(header)
        set_run_font(run, 10.2, bold=True)
    for row_values in rows:
        row = tbl.add_row()
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            set_column_width(cell, widths[i])
            set_cell_border(cell)
            set_cell_margins(cell, top=90, start=120, bottom=90, end=120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            set_paragraph(p, after=0, line=1.12)
            run = p.add_run(value)
            set_run_font(run, 9.7)


def build() -> None:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(title, before=24, after=6)
    run = title.add_run("第一版音乐游戏生成智能体")
    set_run_font(run, 24, bold=True, color=BLUE_DARK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(subtitle, after=12)
    run = subtitle.add_run("赛道二项目汇报稿：大语言模型创新应用开发")
    set_run_font(run, 13.5, bold=True, color=MUTED)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(meta, after=16)
    run = meta.add_run(f"参考赛前培训 PPT 要求整理 | 生成日期：{date.today().isoformat()}")
    set_run_font(run, 10, color=MUTED)

    callout(
        doc,
        "赛道二一句话定位",
        "本项目不是普通的低代码智能体，也不是只会生成文字的教案助手，而是一个把学校大模型平台 API 与音乐教育业务逻辑、音频/MIDI 处理、游戏模板和网页生成系统深度融合的教育创新应用。",
        fill=FILL_BLUE,
        accent=BLUE_DARK,
    )

    heading(doc, "1. 按培训 PPT 提炼出的赛道二汇报重点", 1)
    table(
        doc,
        ["PPT 要求", "赛道二关注点", "本项目应如何回应"],
        [
            ["API 调用与技术融合", "讲清如何对接学校大模型平台 API，并与业务系统或第三方模块结合", "重点讲 ChatECNU/ECNU 大模型作为教案设计大脑，连接 FastAPI、GenerationPipeline、模板工厂、OpenCode、Basic Pitch 与前端游戏运行时。"],
            ["交互体验与生成质量", "展示对话流畅度、内容准确性、逻辑相关性和创造性", "展示教师从教案/一句需求到可运行游戏的流程，强调生成结果不是文本，而是能听、能玩、能判定、能修改的课堂网页。"],
            ["多模态能力展示", "如有图像、语音、音频、视频能力，应作为亮点", "突出曲谱 JPG/PNG/PDF 视觉识别、音频转 MIDI、SoundFont 播放、学生网页交互和可视化游戏场景。"],
            ["应用价值", "紧密结合教育场景，讲清师生收益和校内推广", "聚焦音乐教师备课效率、课堂互动质量、学生音乐理解与创作体验，以及校内音乐课堂/社团/公开课推广。"],
            ["技术价值权重高", "技术性能、系统效率、可用性是主要评价点", "强调结构化合同、模板实例化、GameVariantSpec、多智能体验收与服务器部署，说明稳定性和可复制性。"],
        ],
        [2100, 2700, 4560],
    )

    heading(doc, "2. 背景与真问题：为什么音乐课堂需要这个智能体", 1)
    para(doc, "音乐课堂中的 AI 需求不是“帮老师写一段教案”这么简单。真正耗时的是把教学目标、曲目材料和课堂活动转成学生愿意参与、教师能掌控、现场能运行的互动任务。")
    bullet(doc, "痛点一：音乐材料难转化。曲谱、音频、节奏型、唱名和曲式知识很难直接变成网页互动。", prefix="痛点一：")
    bullet(doc, "痛点二：备课交互成本高。教师常常需要手工找素材、做课件、设计规则、测试链接。", prefix="痛点二：")
    bullet(doc, "痛点三：课堂互动不足。传统 PPT 或口头练习缺少即时反馈，学生体验感和参与度有限。", prefix="痛点三：")
    bullet(doc, "AI 必要性：大模型擅长理解教案语义和生成方案，但必须和音乐处理、游戏模板、网页交付、质量验收结合，才能真正进入课堂。", prefix="AI 必要性：")

    heading(doc, "3. 解决方案：一个大模型驱动的音乐游戏生成系统", 1)
    callout(
        doc,
        "产品定位",
        "教师输入教案、曲目、曲谱、音频或一句课堂需求，系统自动分析教学目标与音乐要素，匹配合适游戏模板或生成 OpenCode 任务包，最终输出可独立打开的音乐课堂网页和小游戏。",
        fill=FILL_GREEN,
        accent=GREEN,
    )
    table(
        doc,
        ["用户输入", "大模型理解", "系统生成", "课堂使用"],
        [
            ["教案文本", "提炼教学目标、课堂环节、音乐要素", "推荐游戏方案、任务规则、教师提示", "课前快速生成互动活动"],
            ["曲谱/图片/PDF", "识别曲目与核心乐句，形成歌曲材料", "绑定节奏、音高、唱名等音乐实体", "让学生围绕真实音乐材料闯关"],
            ["音频材料", "结合音频/MIDI 分析得到可确认候选", "生成听辨、节奏、音色或改编工具", "支持听、比、改、玩"],
            ["一句需求", "判断活动类型和玩法目标", "实例化节拍、节奏、音高、曲式等模板", "学生直接打开网页体验"],
        ],
        [1800, 2600, 2700, 2260],
    )

    heading(doc, "4. 核心创新：赛道二要重点讲的“凭什么是你”", 1)
    bullet(doc, "创新一：大模型 API 不只生成文本，而是生成可执行合同。系统把教师需求转成 activity_type、interaction_model、scoring、runtime_behaviors、visual_theme、GameVariantSpec 等结构化字段。", prefix="创新一：")
    bullet(doc, "创新二：音乐教育业务逻辑深度嵌入。系统内置课程目标、学段特征、音乐要素、模板能力图谱和教案贴合检查，不让大模型凭空编玩法。", prefix="创新二：")
    bullet(doc, "创新三：模板实例化和代码生成双通道。成熟课堂游戏直接走模板，复杂需求再交给 OpenCode 任务包，兼顾稳定性和扩展性。", prefix="创新三：")
    bullet(doc, "创新四：多智能体验收。生成后由 music-logic-agent、lesson-fit-agent、browser-qa-agent、code-interpreter、repair-agent 等检查并修复。", prefix="创新四：")
    bullet(doc, "创新五：多模态音乐材料链路。支持文本教案、曲谱图片/PDF、音频、MIDI、网页交互和游戏视觉资源。", prefix="创新五：")

    heading(doc, "5. 技术实现：从大模型 API 到可运行网页", 1)
    table(
        doc,
        ["技术层", "关键实现", "解决的问题"],
        [
            ["模型接入层", "ModelGateway 对接 ChatECNU/ECNU 大模型 API，DeepSeek/OpenAI/Ollama 可作为备用或本地测试方案", "让大模型承担教案理解、目标提炼、玩法决策和提示词润色"],
            ["生成合同层", "TOOL_SPEC_SCHEMA、lesson_context、lesson_fit、GameVariantSpec、frontend-handoff-contract", "把自然语言转为可验证、可传递、可执行的结构化任务"],
            ["生成编排层", "GenerationPipeline 管理规格生成、活动锁定、素材准备、OpenCode 任务包、执行验收", "避免生成过程散乱，形成稳定工程闭环"],
            ["模板工厂层", "game_template_registry + game_workflow_orchestrator，维护七类成熟音乐游戏模板", "把常见课堂玩法配置化，减少重新生成网页的不确定性"],
            ["音乐处理层", "Basic Pitch、librosa、pretty_midi、SoundFont、FFmpeg、Web Audio", "把音频和 MIDI 转化为课堂可听、可调、可互动的音乐材料"],
            ["前端游戏层", "React + Vite + Radix Themes + Zustand + Phaser 2D", "构建教师控制台与学生端游戏场景，支持键盘、点击、触屏、试听和反馈"],
            ["质量保障层", "execution_orchestrator 多智能体验收与 repair-agent 自动修复", "提高生成产物可用性、音乐合理性和教案贴合度"],
        ],
        [1800, 4300, 3260],
    )

    heading(doc, "6. 核心流程图讲解词", 1)
    for item in [
        "教师输入教案、曲谱、音频或一句需求。",
        "ChatECNU/ECNU 大模型 API 先理解教学目标、音乐要素和课堂环节。",
        "系统把理解结果写入结构化合同，而不是直接让模型写最终页面。",
        "模板能力图谱判断适合节拍、节奏、音高、唱名、音色、曲式还是创编模板。",
        "成熟模板直接实例化；复杂需求生成 OpenCode 任务包，交给代码智能体扩展。",
        "前端生成独立网页，学生可以听、点、拖、唱、闯关并获得反馈。",
        "多智能体检查音乐逻辑、教案贴合、浏览器可用性、代码结构和版本快照。",
    ]:
        number(doc, item)

    heading(doc, "7. 交互体验与生成质量展示", 1)
    table(
        doc,
        ["展示点", "建议现场呈现", "要说出的价值"],
        [
            ["高质量对话", "展示教师输入一句需求或上传教案后，系统给出推荐游戏方案", "说明大模型理解不是泛泛聊天，而是能落到教学目标和音乐要素"],
            ["生成结果", "打开一个节奏复刻、音高爬梯或节拍守卫游戏页面", "说明系统最终交付的是可运行课堂工具，而不是文字建议"],
            ["继续修改", "演示“把难度降低一点/换成三拍子/改成音高练习”", "体现对话式迭代和当前游戏 patch 能力"],
            ["多模态材料", "展示曲谱图片/PDF 或音频材料进入系统后的分析链路", "突出赛道二的技术丰富度和教育场景深度"],
            ["质量报告", "展示执行报告或说明多智能体验收结果", "证明生成质量可检查、可修复、可复用"],
        ],
        [1900, 3600, 3860],
    )

    heading(doc, "8. 应用价值：校内推广怎么讲", 1)
    bullet(doc, "对教师：减少从教案到互动活动的制作成本，提高备课效率和课堂设计质量。", prefix="对教师：")
    bullet(doc, "对学生：把听辨、节奏、音高、曲式和创编变成可操作任务，增强参与感和音乐理解。", prefix="对学生：")
    bullet(doc, "对学校：可用于音乐课堂、社团活动、公开课展示、数字化教学资源建设。", prefix="对学校：")
    bullet(doc, "可复制性：同一套模板和生成链路可迁移到不同曲目、学段、课堂环节和校内教学场景。", prefix="可复制性：")
    bullet(doc, "社会价值：技术服务育人目标，帮助教师更有温度地组织音乐体验，而不是用 AI 替代教师。", prefix="社会价值：")

    heading(doc, "9. 推荐 PPT 结构", 1)
    table(
        doc,
        ["页码", "页面主题", "核心内容"],
        [
            ["1", "封面", "项目名称、赛道二、团队/指导老师、关键词：大模型 API + 音乐教育 + 生成式互动网页"],
            ["2", "背景与痛点", "音乐课堂备课、材料转化、互动反馈的真问题"],
            ["3", "AI 必要性", "为什么普通课件/普通教案生成不够，必须用大模型与音乐技术融合"],
            ["4", "产品定位", "一句话讲清：教师输入需求，系统生成可上课的音乐游戏"],
            ["5", "核心功能", "教案分析、模板匹配、音乐处理、网页生成、继续修改、作品管理"],
            ["6", "赛道二技术架构", "ChatECNU API + FastAPI + GenerationPipeline + 模板工厂 + OpenCode + 前端游戏"],
            ["7", "API 融合细节", "模型如何生成结构化合同，如何和业务逻辑、音乐材料、模板能力连接"],
            ["8", "多模态能力", "教案文本、曲谱图片/PDF、音频转 MIDI、网页交互与游戏视觉"],
            ["9", "生成质量保障", "GameVariantSpec、多智能体验收、音乐逻辑校验、浏览器可用性检查"],
            ["10", "成果展示", "放 2-3 张界面截图或现场演示：控制台、推荐方案、学生端游戏"],
            ["11", "应用价值", "校内音乐课堂推广、教师效率、学生体验、资源复用"],
            ["12", "总结展望", "短期完善材料实体层与模板能力图谱，长期建设音乐教育生成平台"],
        ],
        [900, 2100, 6360],
    )

    heading(doc, "10. 5-8 分钟汇报稿", 1)
    para(doc, "开场：各位评委老师好，我们参加的是赛道二：大语言模型创新应用开发。我们的项目叫“第一版音乐游戏生成智能体”，它面向中小学音乐课堂，目标不是让 AI 只写教案，而是让 AI 帮教师生成真正能在课堂上打开、操作和反馈的音乐游戏。")
    para(doc, "问题：音乐教师在备课时，最难的不是想一个活动名称，而是把曲谱、音频、节奏型、唱名、曲式这些音乐材料转成学生能参与的互动任务。传统 PPT 缺少即时反馈，普通大模型又容易停留在文字建议，无法直接进入课堂。")
    para(doc, "方案：我们的系统把学校大模型平台 API 接入到完整业务链路中。教师输入教案、曲谱、音频或一句需求后，大模型先提炼教学目标和音乐要素，再由生成流水线匹配模板、生成配置、调用音乐处理能力，最后输出独立网页。学生端可以听、点、拖、唱、闯关，教师端可以预览、下载和继续修改。")
    para(doc, "技术：赛道二的关键是技术融合。我们使用 FastAPI 做后端，React/Vite 做控制台，Phaser 2D 做学生端游戏，Basic Pitch 和 pretty_midi 做音频与 MIDI 处理。大模型不直接写最终答案，而是生成结构化合同，例如 GameVariantSpec 和 frontend handoff contract，再交给模板工厂或 OpenCode 执行。")
    para(doc, "质量：为了保证生成结果能上课，我们设计了多智能体验收层，包括音乐逻辑检查、教案贴合检查、浏览器可用性检查、代码结构检查、自动修复和版本记录。也就是说，系统不仅负责生成，还负责判断生成结果是否真的符合音乐课堂。")
    para(doc, "价值：这个项目的价值在于降低音乐教师制作互动活动的门槛，让学生从“看课件、听讲解”变成“边听、边玩、边判断、边创作”。它可以在校内音乐课堂、社团、公开课和数字化资源建设中推广。")
    para(doc, "收束：我们相信，技术是手段，育人是目的。这个智能体希望把大模型能力变成教师可掌控、学生能参与、课堂能落地的音乐学习体验。")

    heading(doc, "11. 答辩高频问题准备", 1)
    table(
        doc,
        ["评委可能问", "建议回答"],
        [
            ["为什么属于赛道二？", "核心创新是大模型 API 与音乐教育业务、音频/MIDI 处理、网页游戏生成和多智能体验收的深度融合，不是单纯低代码流程。"],
            ["大模型具体做了什么？", "负责教案理解、目标提炼、音乐要素识别、玩法决策、提示词润色和结构化规格生成，后续由工程系统执行与校验。"],
            ["如何避免大模型胡编音乐知识？", "通过受控模板、音乐实体、GameVariantSpec、确认门槛、音乐逻辑智能体和教案贴合智能体共同约束。"],
            ["技术壁垒在哪里？", "壁垒在于音乐教育知识、生成合同、模板能力图谱、音乐处理链路、可运行前端游戏和质量验收闭环的组合。"],
            ["目前能落地吗？", "可以作为服务器部署，教师和评委只需打开网页；后台统一完成模型调用、音频处理、网页生成和作品管理。"],
            ["后续怎么迭代？", "短期强化曲谱/音频到音乐实体的解析，中期完善模板能力图谱和对话 patch，长期建设面向音乐教育的可复用生成平台。"],
        ],
        [2600, 6760],
    )

    heading(doc, "12. 最后点题", 1)
    callout(
        doc,
        "汇报结尾可用",
        "本项目的核心不是炫技，而是让大语言模型真正进入音乐教育的业务闭环：理解教案、处理音乐材料、生成互动游戏、检查课堂可用性，最终服务教师教学和学生的音乐理解、表达与创造。",
        fill=FILL_BLUE,
        accent=BLUE_DARK,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
