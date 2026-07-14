from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "展板精简文案.md"
OUTPUT = Path.home() / "Desktop" / "不亦乐乎展板精简文案.docx"


TITLE = "不亦乐乎展板精简文案"


def set_run_font(run, size: int | float, bold: bool = False, color: str = "000000") -> None:
    run.font.name = "PingFang SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before: int = 0, after: int = 6, line: float = 1.15) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def add_title(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(paragraph, after=14)
    run = paragraph.add_run(TITLE)
    set_run_font(run, 20, bold=True, color="17365D")


def add_heading(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, before=10, after=5)
    run = paragraph.add_run(text)
    set_run_font(run, 14, bold=True, color="1F4D78")


def add_body(doc: Document, text: str) -> None:
    paragraph = doc.add_paragraph()
    set_paragraph_spacing(paragraph, after=6)
    run = paragraph.add_run(text)
    set_run_font(run, 11)


def parse_sections(markdown_text: str) -> list[tuple[str, list[str]]]:
    sections: list[tuple[str, list[str]]] = []
    current_heading: str | None = None
    current_lines: list[str] = []

    for raw_line in markdown_text.splitlines():
        line = raw_line.strip()
        if not line:
            continue
        if line.startswith("# "):
            continue
        if line.startswith("## "):
            if current_heading is not None:
                sections.append((current_heading, current_lines))
            current_heading = line[3:].strip()
            current_lines = []
            continue
        current_lines.append(line)

    if current_heading is not None:
        sections.append((current_heading, current_lines))

    return sections


def build() -> Path:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)
    section.left_margin = Inches(0.95)
    section.right_margin = Inches(0.95)

    normal = doc.styles["Normal"]
    normal.font.name = "PingFang SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(11)

    add_title(doc)

    content = SOURCE.read_text(encoding="utf-8")
    for heading, lines in parse_sections(content):
        add_heading(doc, heading)
        for line in lines:
            add_body(doc, line)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
