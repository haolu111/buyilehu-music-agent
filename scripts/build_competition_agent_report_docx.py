from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path.home() / "Desktop" / "第一版音乐游戏生成智能体比赛汇报版.docx"

FONT_LATIN = "Arial"
FONT_CN = "PingFang SC"
INK = "1E293B"
MUTED = "64748B"
BLUE = "1D4ED8"
BLUE_DARK = "1E3A8A"
GREEN = "047857"
AMBER = "B45309"
ROSE = "BE123C"
VIOLET = "6D28D9"
FILL_BLUE = "EFF6FF"
FILL_GREEN = "ECFDF5"
FILL_AMBER = "FFFBEB"
FILL_VIOLET = "F5F3FF"
FILL_ROSE = "FFF1F2"
FILL_SLATE = "F8FAFC"
BORDER = "CBD5E1"
TABLE_HEADER = "E2E8F0"
PAGE_WIDTH_DXA = 9360


def set_run_font(run, size: float, *, bold: bool = False, color: str = INK, italic: bool = False) -> None:
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph(paragraph, *, before: float = 0, after: float = 6, line: float = 1.15) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 100, start: int = 140, bottom: int = 100, end: int = 140) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = PAGE_WIDTH_DXA, indent_dxa: int = 120) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent_dxa))
    tbl_ind.set(qn("w:type"), "dxa")


def set_column_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def mark_row_as_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def add_para(doc: Document, text: str, *, size: float = 10.8, color: str = INK, bold: bool = False, after: float = 6) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, after=after)
    run = p.add_run(text)
    set_run_font(run, size, bold=bold, color=color)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    if level == 1:
        set_paragraph(p, before=16, after=8)
        run = p.add_run(text)
        set_run_font(run, 16, bold=True, color=BLUE_DARK)
    elif level == 2:
        set_paragraph(p, before=10, after=5)
        run = p.add_run(text)
        set_run_font(run, 13, bold=True, color=BLUE)
    else:
        set_paragraph(p, before=8, after=4)
        run = p.add_run(text)
        set_run_font(run, 12, bold=True, color=INK)


def add_bullet(doc: Document, text: str, *, bold_prefix: str = "") -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    set_paragraph(p, after=4, line=1.15)
    if bold_prefix and text.startswith(bold_prefix):
        prefix = p.add_run(bold_prefix)
        set_run_font(prefix, 10.6, bold=True, color=INK)
        rest = p.add_run(text[len(bold_prefix) :])
        set_run_font(rest, 10.6, color=INK)
    else:
        run = p.add_run(text)
        set_run_font(run, 10.6, color=INK)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.35)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    set_paragraph(p, after=4, line=1.15)
    run = p.add_run(text)
    set_run_font(run, 10.6, color=INK)


def add_callout(doc: Document, title: str, body: str, *, fill: str = FILL_BLUE, accent: str = BLUE_DARK) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, PAGE_WIDTH_DXA)
    mark_row_as_header(table.rows[0])
    cell = table.cell(0, 0)
    shade_cell(cell, fill)
    set_cell_border(cell, color="BFDBFE" if fill == FILL_BLUE else BORDER)
    set_cell_margins(cell, top=140, start=180, bottom=140, end=180)
    p = cell.paragraphs[0]
    set_paragraph(p, after=3)
    run = p.add_run(title)
    set_run_font(run, 11.5, bold=True, color=accent)
    p = cell.add_paragraph()
    set_paragraph(p, after=0)
    run = p.add_run(body)
    set_run_font(run, 10.6, color=INK)


def add_two_col_table(doc: Document, rows: list[tuple[str, str]], *, widths: tuple[int, int] = (2300, 7060)) -> None:
    table = doc.add_table(rows=0, cols=2)
    set_table_width(table, PAGE_WIDTH_DXA)
    for label, detail in rows:
        row = table.add_row()
        if len(table.rows) == 1:
            mark_row_as_header(row)
        for i, value in enumerate((label, detail)):
            cell = row.cells[i]
            set_column_width(cell, widths[i])
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            shade_cell(cell, FILL_SLATE if i == 0 else "FFFFFF")
            p = cell.paragraphs[0]
            set_paragraph(p, after=0, line=1.15)
            run = p.add_run(value)
            set_run_font(run, 10.4, bold=(i == 0), color=BLUE_DARK if i == 0 else INK)


def add_matrix_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, PAGE_WIDTH_DXA)
    mark_row_as_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_column_width(cell, widths[i])
        shade_cell(cell, TABLE_HEADER)
        set_cell_border(cell)
        set_cell_margins(cell, top=90, start=120, bottom=90, end=120)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        set_paragraph(p, after=0)
        run = p.add_run(header)
        set_run_font(run, 10.2, bold=True, color=INK)
    for row_values in rows:
        row = table.add_row()
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            set_column_width(cell, widths[i])
            set_cell_border(cell)
            set_cell_margins(cell, top=90, start=120, bottom=90, end=120)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            set_paragraph(p, after=0, line=1.12)
            run = p.add_run(value)
            set_run_font(run, 9.7, color=INK)


def add_flow(doc: Document) -> None:
    steps = [
        ("教师输入", "教案、曲谱、音频或一句课堂需求", FILL_BLUE, BLUE),
        ("教案设计大脑", "提炼教学目标、音乐要素、适合游戏化的环节", FILL_GREEN, GREEN),
        ("生成流水线", "形成结构化规格、玩法蓝图、前端交付合同", FILL_AMBER, AMBER),
        ("模板或 OpenCode", "成熟模板直接实例化，复杂需求生成任务包", FILL_VIOLET, VIOLET),
        ("音乐与页面执行", "MIDI、SoundFont、Phaser、React 共同生成互动网页", FILL_BLUE, BLUE),
        ("多智能体验收", "音乐逻辑、教案贴合、浏览器可用性、版本快照", FILL_ROSE, ROSE),
        ("课堂产物", "可独立打开、预览、下载、继续修改的音乐游戏", FILL_SLATE, BLUE_DARK),
    ]
    table = doc.add_table(rows=len(steps), cols=1)
    set_table_width(table, 8800, indent_dxa=280)
    mark_row_as_header(table.rows[0])
    for idx, (title, body, fill, accent) in enumerate(steps):
        cell = table.rows[idx].cells[0]
        shade_cell(cell, fill)
        set_cell_border(cell, color=BORDER)
        set_cell_margins(cell, top=110, start=160, bottom=110, end=160)
        p = cell.paragraphs[0]
        set_paragraph(p, after=2)
        run = p.add_run(f"{idx + 1}. {title}")
        set_run_font(run, 11.2, bold=True, color=accent)
        p = cell.add_paragraph()
        set_paragraph(p, after=0)
        run = p.add_run(body)
        set_run_font(run, 9.8, color=INK)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.78)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_CN)
    normal.font.size = Pt(10.8)
    normal.font.color.rgb = RGBColor.from_string(INK)

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_paragraph(header, after=0)
    run = header.add_run("不亦乐乎音乐游戏生成智能体 | 比赛汇报版")
    set_run_font(run, 9, color=MUTED)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(footer, after=0)
    run = footer.add_run("面向中小学音乐课堂的互动网页与音乐小游戏生成系统")
    set_run_font(run, 9, color=MUTED)


def build_doc() -> None:
    doc = Document()
    configure_document(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(title, before=28, after=6)
    run = title.add_run("第一版音乐游戏生成智能体")
    set_run_font(run, 24, bold=True, color=BLUE_DARK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(subtitle, after=14)
    run = subtitle.add_run("比赛汇报材料：定位、功能、架构与技术全景介绍")
    set_run_font(run, 13.5, color=MUTED, bold=True)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(meta, after=18)
    run = meta.add_run(f"生成日期：{date.today().isoformat()}    项目目录：第一版")
    set_run_font(run, 10, color=MUTED)

    add_callout(
        doc,
        "一句话定位",
        "这是一个面向中小学音乐课堂的音乐游戏生成智能体。教师输入教案、曲目、曲谱、音频或一句课堂需求，系统自动生成可独立打开、可直接上课使用的互动网页工具与音乐小游戏。",
        fill=FILL_BLUE,
        accent=BLUE_DARK,
    )

    add_heading(doc, "1. 项目定位：它解决什么问题", 1)
    add_para(
        doc,
        "第一版智能体的核心对象不是程序员，而是音乐教师。它把“备课想法”转化为“课堂可运行产物”，让教师不必从零设计交互页面、游戏规则、音乐材料和技术实现。",
    )
    add_bullet(doc, "面向场景：中小学音乐课堂，尤其是需要听辨、表现、创编、闯关和即时反馈的活动。", bold_prefix="面向场景：")
    add_bullet(doc, "使用对象：教师在主控台输入需求，学生使用生成后的独立网页或小游戏。", bold_prefix="使用对象：")
    add_bullet(doc, "产物形态：不是单纯生成教案文本，而是生成能播放、能点击、能拖拽、能判定的课堂网页工具。", bold_prefix="产物形态：")
    add_bullet(doc, "系统边界：主界面只负责需求输入、进度展示和作品管理，真正进入课堂的是生成后的独立页面。", bold_prefix="系统边界：")

    add_heading(doc, "2. 核心功能全景", 1)
    add_matrix_table(
        doc,
        ["功能模块", "教师输入", "系统输出", "课堂价值"],
        [
            ["需求理解", "自然语言需求、学段、曲目、教学目标", "结构化生成规格、活动类型、玩法方向", "降低提示词门槛，把模糊想法变成可执行任务"],
            ["教案分析", "教案文本、上传文件、补充要求", "目标、音乐要素、课堂环节、游戏化切入点", "让生成结果贴合真实教学，而不是泛泛小游戏"],
            ["聆听编辑", "音频文件和音乐要素调整要求", "MIDI 识别、调式/速度/节奏/音色调整工具", "支持听辨、改编、比较和音乐要素体验"],
            ["表现闯关", "曲目、节奏型、学段、评价规则", "阶梯式关卡、学生任务、教师提示、通关规则", "把表现训练变成有目标、有反馈的闯关任务"],
            ["创造拼图", "调式、主音、情绪、小节数、创作方式", "素材卡、拖拽拼图、网格旋律线、双声部试听", "支持学生进行音乐创编和即时试听"],
            ["音乐小游戏", "音高、节奏、唱名、音色、曲式等概念", "节拍守卫、音高爬梯、节奏复刻等可玩模板", "把抽象音乐知识转成可操作规则和反馈"],
            ["作品管理", "生成后的作品和修改要求", "预览、下载、继续修改、版本记录", "方便课前准备、课中使用和课后复用"],
        ],
        [1450, 2450, 2750, 2710],
    )

    add_heading(doc, "3. 总体架构：从输入到课堂网页", 1)
    add_para(doc, "系统采用“前端控制台 + FastAPI 后端 + 大模型教案设计大脑 + 生成流水线 + 模板/代码执行 + 多智能体验收”的分层架构。")
    add_flow(doc)

    add_heading(doc, "4. 技术架构分层", 1)
    add_two_col_table(
        doc,
        [
            ("前端控制台", "React + Vite + Radix Themes + Zustand。负责教师配置、模板选择、参数调整、产物预览和学生端入口。"),
            ("后端服务", "FastAPI。提供登录注册、作品管理、任务队列、上传处理、生成任务、修改任务、部署状态和健康检查等接口。"),
            ("模型大脑层", "优先接入 ChatECNU 作为教案设计大脑，可回退 DeepSeek、OpenAI、Ollama 或本地规则。职责是理解教案、提炼目标、选择玩法和生成规格。"),
            ("生成编排层", "GenerationPipeline 统一管理规格生成、活动锁定、素材准备、OpenCode 任务包、本地产物渲染、执行验收和失败处理。"),
            ("模板工厂层", "game_template_registry 与 workflow_orchestrator 维护成熟模板和配置实例，例如节拍守卫、音高爬梯、节奏复刻、唱名打靶、音色侦探、曲式寻宝、拼图创编。"),
            ("音乐处理层", "Basic Pitch、librosa、pretty_midi、mido、SoundFont、FFmpeg 等能力用于音频转 MIDI、MIDI 事件提取、调式与节奏处理、离线或网页端播放。"),
            ("游戏运行层", "学生端使用 React 与 Phaser 2D 组织游戏场景、状态机、HUD、音频反馈、键盘/点击/触屏输入和通关逻辑。"),
            ("验收修复层", "execution_orchestrator 组织多个执行智能体检查前端文件、音乐逻辑、教案贴合度、浏览器可用性、代码结构、安全修复和版本快照。"),
        ]
    )

    add_heading(doc, "5. 关键工作流", 1)
    add_heading(doc, "5.1 教案生成音乐游戏", 2)
    for item in [
        "教师上传教案或输入课堂片段，系统先解析文本、曲目、音乐要素和目标环节。",
        "教案设计大脑增强分析结果，形成 lesson_context、lesson_fit、lesson_template_contract 等中间合同。",
        "模板路由器根据教学重点选择合适玩法，例如节奏目标匹配 rhythm_echo_core，音高目标匹配 pitch_ladder_core。",
        "GameVariantSpec 记录个性化事实来源，包括材料实体、模板参数、确认门槛和前端执行合同。",
        "前端游戏运行时消费配置，不需要每次重写整张网页；教师继续修改时优先 patch 当前游戏合同。",
    ]:
        add_number(doc, item)

    add_heading(doc, "5.2 直接生成课堂工具", 2)
    for item in [
        "教师直接描述需求，例如“给三年级做一个强弱拍闯关游戏”。",
        "模型网关把需求转为 activity_type、interaction_model、scoring、runtime_behaviors、visual_theme 等规格字段。",
        "系统先判断是否命中成熟模板，命中则直接实例化；不命中则生成 OpenCode 任务包交给代码智能体执行。",
        "产物完成后进入多智能体验收，生成执行报告和版本快照。",
    ]:
        add_number(doc, item)

    add_heading(doc, "6. 子智能体分工", 1)
    add_matrix_table(
        doc,
        ["智能体", "定位", "主要职责"],
        [
            ["music-requirement-planner", "需求规划", "拆解教师自然语言，确定活动类型、教学目标和生成约束。"],
            ["lesson-game-designer", "教案设计", "从教案中提炼目标、音乐要素、可游戏化环节和课堂闭环。"],
            ["frontend-artifact-builder", "前端产物", "检查网页文件、交互入口、资源引用和课堂页面结构。"],
            ["music-logic-agent", "音乐逻辑", "校验调式、节奏、关卡、素材、判定规则是否符合音乐课堂逻辑。"],
            ["lesson-fit-agent", "教案贴合", "检查结果是否贴合课例目标、教学环节、音乐要素和迁移任务。"],
            ["browser-qa-agent", "浏览器可用性", "模拟页面打开、按钮、表单、脚本和资源风险。"],
            ["music-tool-calculator", "可验证指标", "计算 BPM、拍数、小节数、素材数量、关卡数量和评价权重。"],
            ["code-interpreter", "结构检查", "解析 JSON、HTML、JavaScript 和 Python 结构，检查配置一致性。"],
            ["repair-agent", "自动修复", "对可安全修复的问题做最小修复，例如补齐 HTML 结束标签、移动端 viewport、基础控件和音频恢复逻辑。"],
            ["versioning-agent", "版本记录", "为每次生成、修改或修复后的产物建立版本快照和清单。"],
        ],
        [2300, 1800, 5260],
    )

    add_heading(doc, "7. 模板与游戏能力", 1)
    add_para(doc, "第一版已经从“每次重新生成网页”升级到“成熟玩法模板 + 配置实例化”的方向。这样可以提高稳定性，也便于在比赛中展示真正可运行、可复用的课堂游戏能力。")
    add_matrix_table(
        doc,
        ["模板", "音乐重点", "典型课堂玩法"],
        [
            ["beat_guardian_core", "节拍、强拍、弱拍、拍号", "学生在正确拍点点击或充能，训练稳定拍和强弱拍感知。"],
            ["rhythm_echo_core", "节奏、时值、休止、切分", "先听目标节奏，再用点击或身体动作复刻，系统给出准确率反馈。"],
            ["pitch_ladder_core", "音高高低、唱名、旋律走向", "听音后选择更高、一样高或更低，角色沿路线前进并支持唱回确认。"],
            ["solfege_target_core", "唱名听辨、内听、唱回", "听到目标音后打靶或连环靶，训练唱名定位。"],
            ["timbre_detective_core", "音色、乐器家族、听辨证据", "像侦探一样根据声音线索辨认乐器和音色特征。"],
            ["form_treasure_core", "曲式、重复、对比、ABA/回旋", "通过段落排序、寻宝和路线判断理解乐曲结构。"],
            ["composition_puzzle_core", "创编、节奏拼图、旋律拼图", "拖拽素材卡生成乐句，试听并根据规则完成创作。"],
        ],
        [2300, 2500, 4560],
    )

    add_heading(doc, "8. 技术亮点", 1)
    add_bullet(doc, "大模型不是直接写最终答案，而是先生成结构化合同，降低幻觉风险。", bold_prefix="大模型")
    add_bullet(doc, "模板和 OpenCode 双通道并存：常见课堂游戏走稳定模板，开放需求走代码生成任务包。", bold_prefix="模板")
    add_bullet(doc, "GameVariantSpec 把音乐材料、模板能力、前端执行和修改历史放在同一事实源中。", bold_prefix="GameVariantSpec")
    add_bullet(doc, "多智能体验收机制覆盖音乐逻辑、教案贴合度、浏览器可用性、代码结构和版本管理。", bold_prefix="多智能体")
    add_bullet(doc, "音乐处理链路覆盖音频转 MIDI、调式和速度调整、SoundFont 采样播放与网页端真实音频反馈。", bold_prefix="音乐处理")
    add_bullet(doc, "教师二次修改优先 patch 当前游戏配置，避免每次修改都重新生成网页。", bold_prefix="二次修改")

    add_heading(doc, "9. 与传统备课方式的区别", 1)
    add_matrix_table(
        doc,
        ["对比维度", "传统方式", "本智能体方式"],
        [
            ["备课入口", "教师手动查资源、做课件、找小游戏", "输入教案或一句需求，系统自动形成游戏方案"],
            ["音乐材料转化", "音频、谱例、节奏型需要教师手工改编", "系统解析音乐材料并绑定到可执行游戏实体"],
            ["互动形式", "多为 PPT、视频或口头练习", "生成可点击、可听、可判定、可闯关的网页"],
            ["个性化", "每个班级、曲目都要重新制作", "通过模板参数和 GameVariantSpec 快速个性化"],
            ["质量检查", "依赖教师人工检查", "多智能体自动检查音乐逻辑、教案贴合和页面可用性"],
            ["复用方式", "复制旧课件再手动改", "作品可预览、下载、继续修改并保留版本"],
        ],
        [1800, 3600, 3960],
    )

    add_heading(doc, "10. 比赛汇报建议话术", 1)
    add_callout(
        doc,
        "开场 30 秒",
        "我们的智能体面向音乐教师，不是生成一段文字教案，而是把教案、曲谱和课堂需求转化成可直接上课的互动网页和音乐小游戏。它的价值在于把音乐课堂中最耗时的“互动活动设计、音乐材料转化、网页实现和质量检查”串成一条自动化链路。",
        fill=FILL_GREEN,
        accent=GREEN,
    )
    add_heading(doc, "演示主线", 2)
    for item in [
        "先展示教师输入：一句需求或一份教案。",
        "再展示系统分析：教学目标、音乐要素、推荐模板、游戏规则。",
        "然后展示生成结果：一个学生可以直接打开玩的音乐网页。",
        "最后强调验收：系统不仅生成，还会检查音乐逻辑、教案贴合度和页面可用性。",
    ]:
        add_number(doc, item)

    add_heading(doc, "11. 评委可能追问与回答", 1)
    add_matrix_table(
        doc,
        ["问题", "回答要点"],
        [
            ["它和普通大模型生成教案有什么不同？", "普通大模型主要输出文本，本系统输出可运行网页和游戏，并且有音乐处理、模板实例化和多智能体验收。"],
            ["如何保证不乱生成音乐知识？", "系统通过受控模板、音乐实体、GameVariantSpec、确认门槛和 music-logic-agent 校验来约束结果。"],
            ["为什么要用模板而不是完全自由生成？", "课堂工具需要稳定可用。成熟模板保证交互和判定可靠，OpenCode 负责复杂扩展，二者互补。"],
            ["教师修改时是否要重新生成？", "升级目标是优先 patch 当前游戏合同和模板参数，只有超出模板能力时才需要完整再生成。"],
            ["部署后评委需要安装环境吗？", "推荐服务器部署，评委只需打开网址；Basic Pitch、OpenCode、音频处理等都在服务器端运行。"],
        ],
        [3000, 6360],
    )

    add_heading(doc, "12. 当前版本边界与后续方向", 1)
    add_bullet(doc, "当前版本已经具备生成控制台、模板工厂、教案分析、音乐处理、作品管理和多智能体验收等核心能力。", bold_prefix="当前版本")
    add_bullet(doc, "后续升级重点是强化音乐材料实体层，让曲谱、文字谱、MIDI、音频都能更稳定落成可执行音乐实体。", bold_prefix="后续升级")
    add_bullet(doc, "模板能力图谱会继续细化，明确每个模板能改什么、不能改什么，避免把不适配需求硬套到模板上。", bold_prefix="模板能力")
    add_bullet(doc, "GameVariantSpec 将继续成为唯一个性化真相来源，支持对话式修改、确认和版本追踪。", bold_prefix="GameVariantSpec")

    add_callout(
        doc,
        "结尾总结",
        "第一版智能体的核心价值可以概括为：让音乐教师把课堂想法、教案和音乐材料快速转化为可直接上课的互动网页与音乐游戏，并通过多智能体流程保证它不仅“生成出来”，而且“贴合教学、能运行、可继续修改”。",
        fill=FILL_BLUE,
        accent=BLUE_DARK,
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_doc()
