from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUTPUT = Path.home() / "Desktop" / "第一版智能体技术亮点流程图.docx"

FONT = "Arial"
FONT_EAST_ASIA = "PingFang SC"
INK = "1E293B"
MUTED = "64748B"
BLUE = "1D4ED8"
BLUE_DARK = "1E3A8A"
BLUE_LIGHT = "EFF6FF"
GREEN = "047857"
GREEN_LIGHT = "ECFDF5"
AMBER = "B45309"
AMBER_LIGHT = "FFFBEB"
VIOLET = "6D28D9"
VIOLET_LIGHT = "F5F3FF"
ROSE = "BE123C"
ROSE_LIGHT = "FFF1F2"
SLATE_LIGHT = "F8FAFC"
BORDER = "CBD5E1"
TABLE_HEADER = "E2E8F0"


FLOW_STEPS = [
    {
        "stage": "1. 教师输入",
        "title": "把课堂想法说清楚",
        "body": "输入教案、曲目、学段或一句自然语言需求。",
        "highlight": "亮点：低门槛入口，不要求教师会写提示词或代码。",
        "fill": BLUE_LIGHT,
        "accent": BLUE,
    },
    {
        "stage": "2. 教案设计大脑",
        "title": "先理解教学目标",
        "body": "提炼音乐要素、课堂环节、适合游戏化的任务。",
        "highlight": "亮点：ChatECNU 优先，DeepSeek / OpenAI / Ollama 可选，本地规则兜底。",
        "fill": GREEN_LIGHT,
        "accent": GREEN,
    },
    {
        "stage": "3. 生成规格",
        "title": "把需求变成结构化任务",
        "body": "形成玩法、关卡、素材、音频、页面和安全规则。",
        "highlight": "亮点：GenerationPipeline 统一编排，避免只生成一段散乱网页代码。",
        "fill": AMBER_LIGHT,
        "accent": AMBER,
    },
    {
        "stage": "4. 模板 / OpenCode 双通道",
        "title": "成熟玩法直接实例化，复杂需求交给代码智能体",
        "body": "节奏、音高、唱名、音色、曲式等模板可复用；新需求可生成 OpenCode 任务包。",
        "highlight": "亮点：快需求走模板，难需求走生成，兼顾稳定性和扩展性。",
        "fill": VIOLET_LIGHT,
        "accent": VIOLET,
    },
    {
        "stage": "5. 音乐与前端能力层",
        "title": "真正做成可操作的音乐网页",
        "body": "Basic Pitch、MIDI 处理、SoundFont 播放、React/Vite 前端共同完成交互。",
        "highlight": "亮点：不只是文本方案，而是能听、能玩、能拖拽、能闯关的课堂工具。",
        "fill": BLUE_LIGHT,
        "accent": BLUE,
    },
    {
        "stage": "6. 多智能体验收",
        "title": "生成后再检查、修复、留版本",
        "body": "音乐逻辑、教案贴合度、浏览器可用性、HTML/JS/JSON 结构都进入验收。",
        "highlight": "亮点：music-logic-agent、lesson-fit-agent、browser-qa-agent 等协同把关。",
        "fill": ROSE_LIGHT,
        "accent": ROSE,
    },
    {
        "stage": "最终产物",
        "title": "独立课堂网页 / 音乐小游戏",
        "body": "教师可预览、下载、继续修改，学生直接打开使用。",
        "highlight": "核心价值：把音乐课堂想法自动转化为可上课的互动网页。",
        "fill": SLATE_LIGHT,
        "accent": BLUE_DARK,
    },
]

PITCH_LADDER_TECH = [
    ("模板定位", "pitch_ladder_core + pitch_ladder_map_shell", "把“音高高低、旋律走向”固定到音高爬梯玩法，生成物锁定为 PitchLadderGame。"),
    ("游戏引擎", "Phaser 2D 场景", "用 960x540 游戏画布绘制山路、云梯、竹节等路线，角色沿曲线跳跃或滑落。"),
    ("核心规则", "PitchLadderGameSystem 状态机", "状态包含 ready、listening、playing、round_clear、mission_success、mission_failed，负责判定高/低/平、能量扣减和通关。"),
    ("音乐输入", "pitch_rounds / pitch_range / midi_offsets", "把 do、re、mi、sol、la 转成每一轮目标音或路线节点；支持方向判断和旋律路径两种模式。"),
    ("声音技术", "Web Audio + 麦克风音高方向识别", "内置音高用振荡器播放；学生唱回时采集麦克风，估计音高轨迹并判断更高、一样高或更低。"),
    ("交互方式", "键盘、点击、触屏滑动", "Space 听音，方向键/WASD 选择高低平，Enter/V 唱回检测；触屏滑动也能选方向。"),
    ("视觉资源", "Content Manifest 皮肤包", "山路、云梯、竹节三套背景、角色动作条和道具图可切换，规则不变但课堂风格可变。"),
    ("质量保障", "模板保真契约 + 自动测试", "测试会检查 runtime_shell、scene_id、Phaser 引擎、皮肤资源、PitchLadderGame 组件没有被生成过程改坏。"),
]


def set_run_font(run, size: float, *, bold: bool = False, color: str = INK) -> None:
    run.font.name = FONT
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph(paragraph, *, before: float = 0, after: float = 6, line: float = 1.15) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 120, start: int = 160, bottom: int = 120, end: int = 160) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")


def add_cell_text(cell, step: dict[str, str]) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    p = cell.paragraphs[0]
    set_paragraph(p, after=2)
    stage = p.add_run(step["stage"])
    set_run_font(stage, 10, bold=True, color=step["accent"])

    p = cell.add_paragraph()
    set_paragraph(p, after=2)
    title = p.add_run(step["title"])
    set_run_font(title, 13, bold=True, color=INK)

    p = cell.add_paragraph()
    set_paragraph(p, after=2)
    body = p.add_run(step["body"])
    set_run_font(body, 10.5, color=INK)

    p = cell.add_paragraph()
    set_paragraph(p, after=0)
    highlight = p.add_run(step["highlight"])
    set_run_font(highlight, 9.5, bold=True, color=step["accent"])


def add_arrow_row(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, before=0, after=0)
    run = p.add_run("↓")
    set_run_font(run, 13, bold=True, color=MUTED)


def set_column_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def add_pitch_ladder_page(doc: Document) -> None:
    doc.add_page_break()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(title, after=3)
    run = title.add_run("特例：音高爬梯的游戏技术实现")
    set_run_font(run, 20, bold=True, color=BLUE_DARK)

    lead = doc.add_paragraph()
    set_paragraph(lead, after=8)
    run = lead.add_run(
        "音高爬梯不是静态网页，而是一套“音乐规则 + Phaser 游戏场景 + 人声辅助判断”的可复用模板。"
        "教师只改教案目标、音域和皮肤，底层游戏运行时保持稳定。"
    )
    set_run_font(run, 10.8, color=INK)

    flow = doc.add_table(rows=1, cols=1)
    set_table_width(flow, 8200)
    cell = flow.cell(0, 0)
    shade_cell(cell, BLUE_LIGHT)
    set_cell_border(cell, color="BFDBFE")
    set_cell_margins(cell, top=120, start=180, bottom=120, end=180)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, after=0)
    run = p.add_run("教案/需求 → 模板匹配 → 生成 pitch_rounds → Phaser 场景运行 → 听音/选择/唱回 → 状态机判定 → 奖励或重试")
    set_run_font(run, 10.5, bold=True, color=BLUE_DARK)

    heading = doc.add_paragraph()
    set_paragraph(heading, before=12, after=5)
    run = heading.add_run("关键技术拆解")
    set_run_font(run, 14, bold=True, color=BLUE_DARK)

    table = doc.add_table(rows=1, cols=3)
    set_table_width(table, 9360)
    widths = [1700, 2600, 5060]
    headers = ["技术层", "实现方式", "讲给评委听"]
    for index, header in enumerate(headers):
      cell = table.rows[0].cells[index]
      set_column_width(cell, widths[index])
      shade_cell(cell, TABLE_HEADER)
      set_cell_border(cell)
      set_cell_margins(cell, top=90, start=120, bottom=90, end=120)
      p = cell.paragraphs[0]
      set_paragraph(p, after=0)
      run = p.add_run(header)
      set_run_font(run, 9.5, bold=True, color=INK)

    for layer, implementation, plain_text in PITCH_LADDER_TECH:
      row = table.add_row()
      values = [layer, implementation, plain_text]
      for index, value in enumerate(values):
        cell = row.cells[index]
        set_column_width(cell, widths[index])
        shade_cell(cell, "FFFFFF")
        set_cell_border(cell)
        set_cell_margins(cell, top=85, start=120, bottom=85, end=120)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        set_paragraph(p, after=0, line=1.12)
        run = p.add_run(value)
        set_run_font(run, 9.2 if index < 2 else 9.0, bold=index == 0, color=BLUE_DARK if index == 0 else INK)

    summary = doc.add_paragraph()
    set_paragraph(summary, before=10, after=0)
    run = summary.add_run("一句话讲法：音高爬梯把抽象的“音高上行、下行、保持”变成可视化路线，学生听到音后选路线、角色移动，系统再用状态机和声音识别给出即时反馈。")
    set_run_font(run, 10.5, bold=True, color=GREEN)


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    normal = doc.styles["Normal"]
    normal.font.name = FONT
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    normal.font.size = Pt(11)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(title, after=3)
    run = title.add_run("第一版智能体技术亮点流程图")
    set_run_font(run, 22, bold=True, color=BLUE_DARK)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(subtitle, after=10)
    run = subtitle.add_run("一句话看懂：教师输入课堂想法，系统自动生成可上课的音乐互动网页。")
    set_run_font(run, 10.5, color=MUTED)

    for index, step in enumerate(FLOW_STEPS):
        table = doc.add_table(rows=1, cols=1)
        set_table_width(table, 7800)
        cell = table.cell(0, 0)
        shade_cell(cell, step["fill"])
        set_cell_border(cell)
        set_cell_margins(cell)
        add_cell_text(cell, step)
        if index != len(FLOW_STEPS) - 1:
            add_arrow_row(doc)

    add_pitch_ladder_page(doc)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(footer, after=0)
    run = footer.add_run("不亦乐乎 · 音乐游戏生成智能体")
    set_run_font(run, 8.5, color=MUTED)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
