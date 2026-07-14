from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = Path.home() / "Desktop" / "第一版智能体说明.docx"


CONTENT = [
    ("title", "第一版智能体说明"),
    ("h1", "1. 智能体定位"),
    ("p", "面向中小学音乐课堂的音乐游戏生成智能体。"),
    ("p", "教师输入课堂需求、教案或音乐材料后，系统生成可独立打开的课堂网页工具或音乐小游戏。"),
    ("p", "主界面只负责需求输入、进度展示和作品管理；真正给学生使用的是生成后的独立网页。"),
    ("h1", "2. 核心功能"),
    ("b", "需求理解：把教师自然语言转成结构化生成规格。"),
    ("b", "教案分析：提炼教学目标、音乐要素和适合游戏化的课堂环节。"),
    ("b", "聆听编辑：上传音频，识别 MIDI，调整调式、节奏、速度和音色。"),
    ("b", "表现闯关：生成阶梯式关卡、学生任务、教师提示和通关规则。"),
    ("b", "创造拼图：生成音乐素材卡、拖拽拼图和网格旋律线创作工具。"),
    ("b", "音乐小游戏：生成节奏、音高、唱名、音色、曲式等可操作游戏。"),
    ("b", "作品管理：支持预览、下载、继续修改和版本记录。"),
    ("h1", "3. 技术架构"),
    (
        "flow",
        "教师输入 -> 前端控制台 React + Vite -> FastAPI 后端 -> ECNU 大模型 API / 教案设计大脑 -> "
        "GenerationPipeline 生成流水线 -> 模板实例化或 OpenCode 生成 -> 多智能体验收与修复 -> 独立课堂网页产物",
    ),
    ("b", "前端：React、Vite、Radix Themes、Zustand。"),
    ("b", "后端：FastAPI、任务队列、用户与作品管理。"),
    ("b", "模型：ECNU 大模型 API 优先理解教案、教学目标和生成需求，DeepSeek 备用，本地规则兜底。"),
    ("b", "音乐处理：Basic Pitch、pretty_midi、mido、SoundFont 采样播放。"),
    ("b", "生成执行：OpenCode 负责网页代码生成与修改；成熟模板可直接实例化。"),
    ("h1", "4. 子 Agent"),
    ("h2", "规划与生成"),
    ("b", "music-requirement-planner：拆解教师需求。"),
    ("b", "lesson-game-designer：设计教案游戏方案。"),
    ("b", "python-music-coder：处理 MIDI 插件。"),
    ("b", "creative-tool-builder：生成创作交互组件。"),
    ("b", "frontend-artifact-builder：生成网页产物。"),
    ("h2", "验收与修复"),
    ("b", "execution-orchestrator：组织验收流程。"),
    ("b", "music-logic-agent：检查音乐逻辑。"),
    ("b", "lesson-fit-agent：检查是否贴合教案。"),
    ("b", "browser-qa-agent：检查页面可打开、可交互。"),
    ("b", "repair-agent：自动修复安全问题。"),
    ("b", "versioning-agent：生成版本快照。"),
    ("b", "code-interpreter：检查 HTML / JS / JSON 结构。"),
    ("h1", "5. Skill 能力"),
    ("b", "体验性活动生成：聆听、识谱、音乐要素调整。"),
    ("b", "表现性闯关生成：课堂关卡和评价规则。"),
    ("b", "表现性游戏模板库：按教学重点选择模板。"),
    ("b", "创造性音乐拼图：素材卡、拖拽、旋律线。"),
    ("b", "音乐小游戏生成：把音乐概念转成游戏规则。"),
    ("b", "一键网页生成：生成完整课堂网页工具。"),
    ("b", "音乐教育知识库：提供学段、课程标准和曲目建议。"),
    ("b", "课堂图片生成：生成或匹配背景、角色和组件素材。"),
    ("h1", "6. 一句话总结"),
    ("quote", "第一版智能体的核心价值是：把教师的音乐课堂想法，自动转化为可直接上课使用的互动网页和音乐游戏。"),
]


def set_run_font(run, size: int | float, bold: bool = False, color: str = "000000") -> None:
    run.font.name = "PingFang SC"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph_spacing(paragraph, before: int = 0, after: int = 6, line: float = 1.15) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), fill)
    tc_pr.append(shading)


def add_footer(document: Document) -> None:
    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run("不亦乐乎音乐游戏生成智能体")
    set_run_font(run, 9, color="666666")


def build() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "PingFang SC"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(10.5)

    for kind, text in CONTENT:
        if kind == "title":
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            set_paragraph_spacing(p, after=12)
            run = p.add_run(text)
            set_run_font(run, 22, bold=True, color="17365D")
        elif kind == "h1":
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=10, after=5)
            run = p.add_run(text)
            set_run_font(run, 14, bold=True, color="1F4D78")
        elif kind == "h2":
            p = doc.add_paragraph()
            set_paragraph_spacing(p, before=6, after=3)
            run = p.add_run(text)
            set_run_font(run, 11.5, bold=True, color="2E74B5")
        elif kind == "b":
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Inches(0.22)
            p.paragraph_format.first_line_indent = Inches(-0.1)
            set_paragraph_spacing(p, after=3)
            run = p.add_run(text)
            set_run_font(run, 10.5)
        elif kind == "flow":
            table = doc.add_table(rows=1, cols=1)
            table.autofit = True
            cell = table.cell(0, 0)
            set_cell_shading(cell, "F2F6FA")
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, after=0)
            run = p.add_run(text)
            set_run_font(run, 10, bold=True, color="17365D")
        elif kind == "quote":
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.18)
            set_paragraph_spacing(p, before=4, after=8)
            run = p.add_run(text)
            set_run_font(run, 11, bold=True, color="17365D")
        else:
            p = doc.add_paragraph()
            set_paragraph_spacing(p, after=5)
            run = p.add_run(text)
            set_run_font(run, 10.5)

    add_footer(doc)
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build()
