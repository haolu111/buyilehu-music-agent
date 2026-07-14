from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


OUTPUT = Path.home() / "Desktop" / "不亦乐乎音乐游戏生成智能体项目书.docx"
ROOT = Path(__file__).resolve().parents[1]
LOGO = ROOT / "app" / "static" / "assets" / "buyilehu-logo-transparent.png"

TITLE_FONT = "方正小标宋简体"
H1_FONT = "黑体"
H2_FONT = "楷体_GB2312"
BODY_FONT = "仿宋_GB2312"
TABLE_FONT = "宋体"
LATIN_FONT = "Times New Roman"

INK = "000000"
MUTED = "666666"
BLUE = "1F4E79"
LIGHT_BLUE = "D9EAF7"
LIGHT_GRAY = "F2F2F2"
BORDER = "999999"


def set_run_font(
    run,
    size: float,
    *,
    cn_font: str = BODY_FONT,
    latin_font: str = LATIN_FONT,
    bold: bool = False,
    color: str = INK,
) -> None:
    run.font.name = latin_font
    run._element.rPr.rFonts.set(qn("w:ascii"), latin_font)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), latin_font)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), cn_font)
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_para(
    paragraph,
    *,
    before: float = 0,
    after: float = 0,
    line_pt: float = 28,
    first_line: bool = False,
    align: WD_ALIGN_PARAGRAPH | None = None,
) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(line_pt)
    if first_line:
        fmt.first_line_indent = Cm(0.74)
    if align is not None:
        paragraph.alignment = align


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shading = tc_pr.find(qn("w:shd"))
    if shading is None:
        shading = OxmlElement("w:shd")
        tc_pr.append(shading)
    shading.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        element = borders.find(qn(f"w:{edge}"))
        if element is None:
            element = OxmlElement(f"w:{edge}")
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 80, start: int = 100, bottom: int = 80, end: int = 100) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def mark_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = tr_pr.find(qn("w:tblHeader"))
    if header is None:
        header = OxmlElement("w:tblHeader")
        tr_pr.append(header)
    header.set(qn("w:val"), "true")


def add_field(paragraph, field_code: str) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(fld_end)


def configure_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(3.7)
    section.bottom_margin = Cm(2.8)
    section.left_margin = Cm(2.8)
    section.right_margin = Cm(2.6)
    section.header_distance = Cm(1.5)
    section.footer_distance = Cm(1.5)

    normal = doc.styles["Normal"]
    normal.font.name = LATIN_FONT
    normal._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
    normal.font.size = Pt(16)

    header = section.header.paragraphs[0]
    set_para(header, line_pt=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = header.add_run("不亦乐乎音乐游戏生成智能体项目书")
    set_run_font(run, 9, cn_font=TABLE_FONT, color=MUTED)

    footer = section.footer.paragraphs[0]
    set_para(footer, line_pt=16, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = footer.add_run("第 ")
    set_run_font(run, 9, cn_font=TABLE_FONT, color=MUTED)
    add_field(footer, "PAGE")
    run = footer.add_run(" 页")
    set_run_font(run, 9, cn_font=TABLE_FONT, color=MUTED)


def add_title(doc: Document, text: str, size: float = 22, after: float = 12) -> None:
    p = doc.add_paragraph()
    set_para(p, after=after, line_pt=32, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run(text)
    set_run_font(run, size, cn_font=TITLE_FONT, bold=True)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    if level == 1:
        set_para(p, before=12, after=6, line_pt=28)
        run = p.add_run(text)
        set_run_font(run, 16, cn_font=H1_FONT, bold=True)
    elif level == 2:
        set_para(p, before=8, after=4, line_pt=28)
        run = p.add_run(text)
        set_run_font(run, 16, cn_font=H2_FONT, bold=True)
    else:
        set_para(p, before=6, after=2, line_pt=26)
        run = p.add_run(text)
        set_run_font(run, 14, cn_font=H1_FONT, bold=True)


def add_body(doc: Document, text: str, *, first_line: bool = True) -> None:
    p = doc.add_paragraph()
    set_para(p, line_pt=28, first_line=first_line)
    run = p.add_run(text)
    set_run_font(run, 16, cn_font=BODY_FONT)


def add_note(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_para(p, line_pt=24, first_line=True)
    run = p.add_run(text)
    set_run_font(run, 12, cn_font=TABLE_FONT, color=MUTED)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    mark_header(table.rows[0])
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        if widths:
            set_cell_width(cell, widths[i])
        shade_cell(cell, LIGHT_GRAY)
        set_cell_border(cell)
        set_cell_margins(cell)
        p = cell.paragraphs[0]
        set_para(p, line_pt=18, align=WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run(header)
        set_run_font(run, 10.5, cn_font=H1_FONT, bold=True)
    for row_values in rows:
        row = table.add_row()
        for i, value in enumerate(row_values):
            cell = row.cells[i]
            if widths:
                set_cell_width(cell, widths[i])
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            set_para(p, line_pt=18)
            run = p.add_run(value)
            set_run_font(run, 10.5, cn_font=TABLE_FONT)
    doc.add_paragraph()


def add_cover(doc: Document) -> None:
    if LOGO.exists():
        p = doc.add_paragraph()
        set_para(p, before=4, after=18, line_pt=18, align=WD_ALIGN_PARAGRAPH.CENTER)
        run = p.add_run()
        run.add_picture(str(LOGO), width=Cm(3.2))

    p = doc.add_paragraph()
    set_para(p, before=38, after=10, line_pt=38, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run("不亦乐乎音乐游戏生成智能体")
    set_run_font(run, 24, cn_font=TITLE_FONT, bold=True)

    p = doc.add_paragraph()
    set_para(p, after=42, line_pt=34, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run("项目书")
    set_run_font(run, 28, cn_font=TITLE_FONT, bold=True)

    rows = [
        ["项目名称", "不亦乐乎音乐游戏生成智能体"],
        ["项目类型", "教育数字化工具 / 人工智能应用 / 音乐课堂互动生成系统"],
        ["建设阶段", "第一版原型已完成，拟进入完善、试点与部署阶段"],
        ["建设单位", "待填写"],
        ["项目负责人", "待填写"],
        ["编制日期", f"{date.today().year}年{date.today().month}月{date.today().day}日"],
    ]
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    for label, value in rows:
        row = table.add_row()
        for i, text in enumerate((label, value)):
            cell = row.cells[i]
            set_cell_width(cell, 2500 if i == 0 else 6500)
            set_cell_border(cell)
            set_cell_margins(cell, top=130, bottom=130)
            if i == 0:
                shade_cell(cell, LIGHT_GRAY)
            p = cell.paragraphs[0]
            set_para(p, line_pt=22, align=WD_ALIGN_PARAGRAPH.CENTER if i == 0 else None)
            run = p.add_run(text)
            set_run_font(run, 12, cn_font=H1_FONT if i == 0 else TABLE_FONT, bold=i == 0)

    p = doc.add_paragraph()
    set_para(p, before=28, line_pt=22, align=WD_ALIGN_PARAGRAPH.CENTER)
    run = p.add_run("项目目录：/Users/shishangbo/codex/第一版")
    set_run_font(run, 10.5, cn_font=TABLE_FONT, color=MUTED)
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    add_title(doc, "目录", 20, after=18)
    entries = [
        "一、项目概况",
        "二、编制依据与格式说明",
        "三、项目背景与建设必要性",
        "四、建设目标与建设原则",
        "五、建设内容与功能范围",
        "六、技术方案与系统架构",
        "七、实施进度计划",
        "八、组织实施与保障措施",
        "九、经费预算与资金安排",
        "十、风险分析与控制措施",
        "十一、预期效益与推广应用",
        "十二、可行性分析与结论",
        "附件一：资料来源与参考依据",
    ]
    for entry in entries:
        p = doc.add_paragraph()
        set_para(p, line_pt=26)
        run = p.add_run(entry)
        set_run_font(run, 14, cn_font=TABLE_FONT)
    doc.add_page_break()


def build_doc() -> None:
    doc = Document()
    configure_document(doc)
    add_cover(doc)
    add_toc(doc)

    add_heading(doc, "一、项目概况", 1)
    add_table(
        doc,
        ["项目要素", "内容"],
        [
            ["项目名称", "不亦乐乎音乐游戏生成智能体"],
            ["项目定位", "面向中小学音乐课堂的音乐游戏生成智能体。教师输入课堂需求、教案或音乐材料后，系统生成可独立打开的课堂网页工具或音乐小游戏。"],
            ["服务对象", "中小学音乐教师、学生、教研人员及音乐教育数字资源建设人员。"],
            ["建设目标", "形成一个可理解教案、可生成互动网页、可复用音乐游戏模板、可进行质量验收和版本管理的课堂智能体系统。"],
            ["当前阶段", "第一版已形成前后端代码、模板工厂、音乐处理链路、作品管理、部署文档和多类测试用例，适合进入完善、试点和公开部署准备阶段。"],
            ["交付形态", "教师端生成控制台、学生端独立网页产物、音乐游戏模板库、OpenCode 任务包、服务器部署包、项目文档与验收测试。"],
        ],
        [2100, 7200],
    )
    add_body(
        doc,
        "本项目以音乐课堂真实教学需求为入口，围绕“教师输入、系统理解、智能生成、课堂交付、继续修改”的完整链路建设。第一版系统的主界面承担需求输入、模块选择、进度展示和作品管理职责，真正进入课堂的是系统生成后的独立网页工具或音乐小游戏。",
    )
    add_body(
        doc,
        "项目的核心价值是把教师的音乐课堂想法、教案文本、曲谱、音频或一句自然语言需求，转化为可点击、可试听、可拖拽、可判定、可复用的课堂互动产物，从而降低音乐教师制作互动资源的技术门槛。"
    )

    add_heading(doc, "二、编制依据与格式说明", 1)
    add_body(
        doc,
        "本项目书按照项目建议书和可行性研究报告的常见结构编制，重点说明项目背景、建设必要性、建设目标、建设内容、技术方案、实施进度、经费预算、风险控制、预期效益和可行性结论。"
    )
    add_table(
        doc,
        ["依据类别", "采用要点", "本项目对应章节"],
        [
            ["项目建议书常见内容", "项目建设必要性和依据、建设内容与规模、投资估算与资金来源、建设进度、经济和社会效益等。", "项目背景、建设内容、预算、进度、效益。"],
            ["可行性研究报告思路", "围绕必要性、可行性、风险可控性，分析需求可靠性、方案可行性、运营有效性、投入合理性和风险可控性。", "需求分析、技术方案、保障措施、风险控制、结论。"],
            ["正式文档排版参考", "采用 A4 纸，页边距按公文式版面设置；标题、正文、表格采用中文正式文档常用字体和固定行距。", "全篇 Word 排版。"],
        ],
        [1900, 5000, 2400],
    )
    add_note(
        doc,
        "格式说明：本文件正文采用仿宋_GB2312三号、固定行距28磅；一级标题采用黑体三号；二级标题采用楷体_GB2312三号加粗；表格采用宋体五号；页面为 A4，页边距为上3.7cm、下2.8cm、左2.8cm、右2.6cm。若提交单位有专门模板，可在不改变内容结构的前提下替换封面和页眉页脚。"
    )

    add_heading(doc, "三、项目背景与建设必要性", 1)
    add_heading(doc, "（一）政策与课程背景", 2)
    add_body(
        doc,
        "依据国家教育数字化战略行动部署及教育部等九部门《关于加快推进教育数字化的意见》，基础教育正处在以数字化赋能教学方式、学习方式和资源供给方式变革的重要阶段。人工智能与教育深度融合，为课堂教学提供了智能化、交互式、可生成的新型工具基础。"
    )
    add_body(
        doc,
        "《义务教育艺术课程标准（2022年版）》强调音乐学科应走向实践化、游戏化、数字化教学路径，并落实以美育人、以美化人、以美培元的根本任务。课程标准提出的审美感知、艺术表现、文化理解、创意实践四个核心素养，要求音乐课堂不仅讲授知识，更要让学生在真实聆听、表现、创编和交流中形成音乐经验。"
    )
    add_body(
        doc,
        "因此，音乐学科数字化不能停留在课件展示或资源播放层面，而应进一步形成能够支持学生参与、音乐感知、课堂反馈和创意实践的互动工具。本项目正是在这一背景下，探索以 AI 音乐课堂游戏生成智能体服务中小学音乐课堂的新路径。"
    )
    add_heading(doc, "（二）现实痛点", 2)
    add_table(
        doc,
        ["痛点", "具体表现", "项目应对方式"],
        [
            ["学生参与弱", "音乐课堂中仍存在机械聆听多、互动任务少的问题，学生容易停留在被动听、被动看、被动回答的状态。", "把聆听、表现、创编任务转化为可点击、可拖拽、可闯关的互动网页，提高课堂参与度。"],
            ["音乐感知难", "节拍、节奏、音高、唱名、音色、曲式等音乐要素相对抽象，学生稳定感知和迁移应用存在困难。", "将音乐要素转化为游戏规则、听辨任务、即时反馈和可视化路径，帮助学生在操作中建立感知。"],
            ["课堂反馈慢", "班级学生人数较多，个体差异明显，教师难以及时了解每位学生的听辨、表现和创编状态。", "通过游戏化任务、关卡反馈、评分规则和作品记录，为教师提供更及时的课堂观察依据。"],
        ],
        [1900, 4300, 3100],
    )
    add_heading(doc, "（三）建设必要性", 2)
    add_body(
        doc,
        "建设本项目有助于回应音乐课堂“学生参与弱、音乐感知难、课堂反馈慢”的现实问题，推动音乐数字资源从“静态展示”转向“可操作、可感知、可反馈、可复用”的互动生成形态。项目以教师语言、教案和音乐材料为入口，降低编程和交互设计门槛，使音乐教师能够围绕本课重点快速生成课堂工具。"
    )
    add_body(
        doc,
        "项目同时具备教学、技术和应用价值：教学上服务审美感知、艺术表现、文化理解、创意实践等核心素养落地；技术上形成教育场景下受控生成、模板实例化和多智能体验收机制；应用上可为比赛展示、校内试点、教研共创和后续产品化提供基础。"
    )

    add_heading(doc, "四、建设目标与建设原则", 1)
    add_heading(doc, "（一）总体目标", 2)
    add_body(
        doc,
        "项目拟建设一个面向中小学音乐课堂的智能体系统，实现从教案和音乐材料到课堂互动网页、音乐小游戏和继续修改能力的闭环，形成可演示、可试点、可部署、可迭代的第一版产品。"
    )
    add_heading(doc, "（二）具体目标", 2)
    add_table(
        doc,
        ["目标类别", "目标内容", "验收要点"],
        [
            ["需求理解目标", "支持教师以自然语言、教案文本、曲谱、音频或音乐材料说明发起生成。", "输出结构化生成规格、教学目标、音乐要素和活动类型。"],
            ["课堂生成目标", "生成聆听编辑、表现闯关、创造拼图和各类音乐小游戏。", "产物可独立打开、可交互、可试听、可下载。"],
            ["模板复用目标", "形成成熟游戏模板库，通过配置实例化完成稳定生成。", "常见节奏、节拍、音高、唱名、音色、曲式、创编需求可命中模板。"],
            ["音乐可信目标", "避免大模型凭空发明不可验证音乐答案。", "通过材料实体、确认门槛、音乐逻辑检查和模板能力边界约束结果。"],
            ["部署运行目标", "支持本地开发和服务器部署两种方式。", "评审或试点用户只需浏览器访问公开网址。"],
        ],
        [1900, 5200, 2200],
    )
    add_heading(doc, "（三）建设原则", 2)
    for item in [
        "教学优先原则：生成结果必须服务真实教学目标，避免只追求炫技效果。",
        "受控生成原则：大模型负责理解和规划，关键音乐答案、模板边界和运行规则必须有结构化合同约束。",
        "模板优先原则：成熟课堂玩法优先走模板实例化，提高稳定性和可复用性。",
        "可验证原则：每个生成产物都应经过音乐逻辑、教案贴合、浏览器可用性和结构一致性检查。",
        "可迭代原则：教师修改优先作用于当前游戏配置和合同，减少完全重生成带来的不确定性。",
    ]:
        add_body(doc, item)

    add_heading(doc, "五、建设内容与功能范围", 1)
    add_heading(doc, "（一）教师端生成控制台", 2)
    add_body(
        doc,
        "教师端生成控制台用于输入课堂需求、上传教案和材料、选择模块、查看生成进度、预览产物、下载资源和继续修改。第一版前端采用 React + Vite + Radix Themes + Zustand 技术栈，后端由 FastAPI 提供接口和任务状态。"
    )
    add_heading(doc, "（二）三类活动母版", 2)
    add_table(
        doc,
        ["活动母版", "主要能力", "适用课堂场景"],
        [
            ["聆听编辑页", "上传音频、Basic Pitch 转 MIDI、调式调性调整、节奏密度和 BPM 调整、音色切换。", "听辨、比较、改编、音乐要素体验。"],
            ["表现关卡页", "根据曲目、学段和目标生成阶梯式闯关、学生任务、教师提示和通关规则。", "节奏模仿、旋律接龙、小组表演、合奏训练。"],
            ["创造拼图页", "根据调式、主音、情绪和小节数生成素材卡，支持拖拽拼图和网格旋律线试听。", "节奏创编、旋律创编、短句拼接、双声部体验。"],
        ],
        [1900, 5200, 2200],
    )
    add_heading(doc, "（三）音乐小游戏模板库", 2)
    add_table(
        doc,
        ["模板名称", "音乐重点", "课堂玩法"],
        [
            ["节拍守卫", "节拍、强弱拍、拍号", "学生在正确拍点点击或完成守卫动作，训练稳定拍和强弱拍感知。"],
            ["节奏复刻", "节奏、时值、休止、切分", "先听示范节奏，再点击或拍击复刻，系统反馈准确率。"],
            ["音高爬梯", "音高高低、旋律走向", "听目标音或音列后判断高低和路线，完成角色前进。"],
            ["唱名打靶", "唱名听辨、内听、模唱", "听目标音后击中对应唱名靶，并支持唱回确认。"],
            ["音色侦探", "乐器音色、音色证据", "听声音线索后匹配乐器或音色特征，并说明依据。"],
            ["曲式寻宝", "重复、对比、ABA、回旋", "通过片段排序和路线判断理解乐曲结构。"],
            ["创编拼图", "节奏、音高、乐句创作", "拖拽素材卡拼出音乐短句，试听后按规则完成创作。"],
        ],
        [1800, 2400, 5100],
    )
    add_heading(doc, "（四）作品管理与二次修改", 2)
    add_body(
        doc,
        "系统支持生成产物预览、下载、删除、版本快照、对话式修改和任务队列查询。教师后续提出“降低难度”“把第2关改成节奏模仿”“替换为某一唱名”等要求时，系统优先修改当前游戏合同和模板参数。"
    )

    add_heading(doc, "六、技术方案与系统架构", 1)
    add_heading(doc, "（一）总体架构", 2)
    add_body(
        doc,
        "系统采用“教师输入 -> 前端控制台 -> FastAPI 后端 -> 教案设计大脑 -> GenerationPipeline 生成流水线 -> 模板实例化或 OpenCode 生成 -> 多智能体验收与修复 -> 独立课堂网页产物”的总体架构。"
    )
    add_table(
        doc,
        ["层级", "技术组成", "职责说明"],
        [
            ["前端层", "React、Vite、Radix Themes、Lucide React、Zustand、Phaser", "提供教师控制台、学生游戏运行时、模板预览、交互输入和音频反馈。"],
            ["接口层", "FastAPI、StaticFiles、任务队列、认证服务", "提供登录注册、文件上传、生成任务、作品管理、部署状态和健康检查。"],
            ["模型层", "ChatECNU 优先，DeepSeek 备用，OpenAI/Ollama 可选，本地规则兜底", "理解教案、提炼教学目标、选择玩法、生成规格和对话修改建议。"],
            ["音乐处理层", "Basic Pitch、pretty_midi、mido、SoundFont、FFmpeg、Tesseract/Audiveris 可选", "完成音频转 MIDI、曲谱或文字材料解析、节奏和音高实体提取、音频渲染。"],
            ["生成执行层", "GenerationPipeline、GameVariantSpec、模板注册表、OpenCode 任务包", "组织规格生成、模板匹配、代码生成、参数注入、运行合同和交付状态。"],
            ["验收层", "music-logic-agent、lesson-fit-agent、browser-qa-agent、repair-agent、versioning-agent", "检查音乐逻辑、教案贴合、页面可用性、结构一致性、修复和版本记录。"],
        ],
        [1700, 3500, 4100],
    )
    add_heading(doc, "（二）关键技术路线", 2)
    add_body(
        doc,
        "项目不把大模型输出直接作为最终网页，而是采用结构化合同和模板工厂控制生成过程。大模型用于理解和规划，模板用于稳定执行，音乐处理工具用于提供可验证材料，多智能体验收用于降低课堂交付风险。"
    )
    add_body(
        doc,
        "对于成熟游戏，系统通过 `GET /api/game-templates`、`GET /api/game-templates/{template_id}` 和 `POST /api/game-instances` 等接口查看模板、获取配置并生成实例；对于非模板化需求，系统生成 `opencode-task.json`、产物规格、配置文件和代码占位文件，交由 OpenCode 执行或后续扩展。"
    )
    add_heading(doc, "（三）生成与验收流程", 2)
    add_table(
        doc,
        ["步骤", "处理内容", "输出物"],
        [
            ["1. 输入采集", "教师输入需求、上传教案、曲谱、音频、MIDI 或 MusicXML。", "原始需求和材料文件。"],
            ["2. 教案理解", "教案设计大脑提炼教学目标、学段、音乐要素、课堂环节和学生行为。", "Lesson Contract、lesson_analysis。"],
            ["3. 模板匹配", "根据音乐重点和活动类型选择成熟模板或开放生成路径。", "template_match、template_decision。"],
            ["4. 参数注入", "将音乐实体、关卡数量、难度、视觉主题和评价规则写入配置。", "GameVariantSpec、instance.config。"],
            ["5. 产物渲染", "生成学生端独立网页和必要资源。", "index.html、runtime-state、素材清单。"],
            ["6. 质量验收", "检查音乐逻辑、教案贴合、浏览器可用性、HTML/JS/JSON 结构和版本记录。", "验收报告、修复记录、版本快照。"],
            ["7. 课堂交付", "教师预览、下载、分享或继续修改。", "可直接上课使用的互动网页或游戏。"],
        ],
        [1100, 5700, 2500],
    )

    add_heading(doc, "七、实施进度计划", 1)
    add_body(
        doc,
        "结合第一版已有代码基础，建议将后续建设划分为六个阶段推进。各阶段可根据比赛、校内试点或产品化节奏进行压缩或延展。"
    )
    add_table(
        doc,
        ["阶段", "周期", "主要任务", "阶段成果"],
        [
            ["第一阶段：现状固化", "第1个月", "梳理第一版能力，固化项目文档、演示流程、部署说明和验收清单。", "项目书、技术手册、演示脚本、部署预检。"],
            ["第二阶段：材料实体层增强", "第2个月", "完善文字谱、MIDI、MusicXML、音频候选到音乐实体的解析和确认流程。", "Lesson Contract、音乐实体测试、确认卡。"],
            ["第三阶段：模板能力图谱", "第3个月", "为成熟模板声明可写字段、不可改骨架、适配范围和拒绝原因。", "template_capability_v2、模板说明。"],
            ["第四阶段：GameVariantSpec 统一", "第4个月", "将教案、材料、模板匹配、执行计划、验收结果和修改历史统一到同一事实源。", "GameVariantSpec v2、兼容测试。"],
            ["第五阶段：模板执行深化", "第5个月", "让教案材料真正影响节奏、节拍、音高、唱名、音色、曲式和创编玩法判定。", "七类模板端到端执行测试。"],
            ["第六阶段：试点部署与优化", "第6个月", "完成服务器部署、教师试用、课堂反馈收集、缺陷修复和版本迭代。", "试点报告、上线包、操作手册。"],
        ],
        [1550, 1350, 4700, 3700],
    )

    add_heading(doc, "八、组织实施与保障措施", 1)
    add_heading(doc, "（一）组织分工建议", 2)
    add_table(
        doc,
        ["角色", "主要职责"],
        [
            ["项目负责人", "统筹项目目标、进度、验收、资源协调和外部汇报。"],
            ["教学设计负责人", "提供音乐课例、课程标准依据、教学目标审校和课堂试用反馈。"],
            ["后端与模型负责人", "维护 FastAPI、模型网关、任务队列、材料解析和部署能力。"],
            ["前端与游戏负责人", "维护教师控制台、学生端游戏运行时、模板交互和移动端体验。"],
            ["音乐技术负责人", "负责 MIDI、音频、SoundFont、节奏/音高/音色等音乐逻辑可靠性。"],
            ["测试与运维负责人", "负责自动化测试、浏览器验收、服务器部署、日志监控和数据安全。"],
        ],
        [2300, 7000],
    )
    add_heading(doc, "（二）质量保障措施", 2)
    for item in [
        "建立需求、设计、生成、验收、修复和版本记录的闭环流程。",
        "对关键合同字段、模板配置、音乐实体解析和前端运行时建立自动化测试。",
        "公开部署前运行服务器预检脚本，确认账号注册、SMTP、Basic Pitch、OpenCode、作品存储等能力可用。",
        "对课堂产物执行浏览器打开、资源引用、脚本执行、音频播放和移动端交互检查。",
        "对音乐答案设置置信度和教师确认机制，避免低置信材料直接成为正式答案。",
    ]:
        add_body(doc, item)

    add_heading(doc, "九、经费预算与资金安排", 1)
    add_body(
        doc,
        "本预算为后续完善、试点和部署阶段的建议预算，具体金额应根据申报单位要求、人员投入方式、服务器规格、模型调用规模和素材采购范围调整。第一版原型已完成的既有投入未在本表中重复计列。"
    )
    add_table(
        doc,
        ["费用类别", "预算金额（元）", "用途说明"],
        [
            ["研发与测试", "45,000", "模板能力图谱、GameVariantSpec、音乐材料解析、端到端测试和缺陷修复。"],
            ["模型调用与算力", "12,000", "ChatECNU、备用模型、音频识别、视觉识谱和试点期间的调用成本。"],
            ["服务器与运维", "10,000", "云服务器、对象存储、HTTPS、日志、备份和部署环境维护。"],
            ["素材与版权", "8,000", "课堂图片、角色素材、音色采样、字体和音乐教学资源合规使用。"],
            ["教学试点与培训", "10,000", "教师试用、培训材料、课堂观察、反馈整理和教研活动。"],
            ["文档与展示", "5,000", "项目书、技术手册、操作手册、比赛展示材料和演示视频。"],
            ["预备费", "10,000", "用于处理不可预见的部署、兼容性和资源采购支出。"],
            ["合计", "100,000", "建议预算合计，最终以实际申报口径为准。"],
        ],
        [2500, 2100, 4700],
    )
    add_note(
        doc,
        "说明：如项目仅用于比赛展示，可压缩服务器、素材、培训和预备费；如进入多校试点，应增加运维、数据安全、教师培训和课堂支持预算。"
    )

    add_heading(doc, "十、风险分析与控制措施", 1)
    add_table(
        doc,
        ["风险类别", "风险表现", "控制措施"],
        [
            ["教学贴合风险", "生成游戏好玩但偏离本课教学目标。", "引入 Lesson Alignment Gate 和教学设计负责人复核。"],
            ["音乐准确性风险", "节奏、音高、唱名、音色等答案来源不清或置信度不足。", "建立音乐实体来源、置信度、教师确认和 music-logic-agent 校验。"],
            ["模型幻觉风险", "大模型编造教学依据或音乐材料。", "采用结构化合同、模板边界、低置信确认卡和本地规则兜底。"],
            ["运行稳定风险", "网页、音频、资源引用或移动端交互异常。", "使用成熟模板、浏览器 QA、静态资源预检和版本快照。"],
            ["部署依赖风险", "OpenCode、Basic Pitch、识谱工具、SMTP 等服务器依赖缺失。", "使用 Docker、预检脚本、状态接口和部署清单。"],
            ["数据安全风险", "上传教案、邮箱账号和学生相关数据管理不当。", "限制采集范围、配置认证、保护 SMTP 密钥、采用 HTTPS 和访问控制。"],
            ["版权合规风险", "音频、图片、字体和素材来源不明。", "优先使用自有、授权或开源素材，建立素材清单和版权说明。"],
        ],
        [1900, 4200, 3200],
    )

    add_heading(doc, "十一、预期效益与推广应用", 1)
    add_heading(doc, "（一）教学效益", 2)
    add_body(
        doc,
        "项目能够帮助音乐教师把课堂重点转化为学生可操作的互动任务，增强听辨、表现、创编和评价环节的参与度。学生通过点击、拖拽、试听、模唱和闯关获得即时反馈，有助于把抽象音乐概念落实为具体音乐行为。"
    )
    add_heading(doc, "（二）技术效益", 2)
    add_body(
        doc,
        "项目形成了面向教育场景的受控生成范式：大模型负责理解和规划，模板负责稳定执行，音乐处理工具负责可验证材料，多智能体负责质量验收。这一技术路线可扩展至其他艺术课程、科学实验和语言学习类互动资源生成。"
    )
    add_heading(doc, "（三）推广应用", 2)
    add_body(
        doc,
        "项目可先在校内音乐教研活动和比赛展示中使用，再选择若干典型课例开展试点，包括节奏课、唱名课、音色听辨课、曲式欣赏课和创编课。试点成熟后，可进一步沉淀为模板库、课例库和教师共创平台。"
    )

    add_heading(doc, "十二、可行性分析与结论", 1)
    add_table(
        doc,
        ["可行性维度", "分析结论"],
        [
            ["技术可行性", "第一版已具备 FastAPI 后端、React 前端、模板工厂、音乐处理、作品管理、部署脚本和大量测试，具备继续完善基础。"],
            ["教学可行性", "系统面向中小学音乐课堂常见活动，覆盖节拍、节奏、音高、唱名、音色、曲式和创编等核心音乐要素。"],
            ["实施可行性", "可按六个月阶段计划推进，优先完善已存在的模板与合同，不依赖一次性大规模重构。"],
            ["运营可行性", "支持服务器部署，终端用户仅需浏览器访问，可降低评委、教师和学生的使用门槛。"],
            ["风险可控性", "项目已设计模型兜底、教师确认、模板边界、多智能体验收、部署预检和版本记录等控制措施。"],
        ],
        [2200, 7100],
    )
    add_body(
        doc,
        "综上，本项目建设目标明确，已有代码和文档基础较完整，技术路线与音乐课堂需求具有匹配性，后续完善和试点具有可行性。建议以第一版原型为基础，继续强化音乐材料实体层、模板能力图谱、GameVariantSpec 统一合同、模板执行深度和服务器部署能力，逐步形成可展示、可试用、可推广的音乐课堂智能体项目。"
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    add_heading(doc, "附件一：资料来源与参考依据", 1)
    add_table(
        doc,
        ["序号", "资料名称", "来源链接", "采用内容"],
        [
            ["1", "《投资项目可行性研究报告编写大纲及说明》发布解读", "https://www.ndrc.gov.cn/xxgk/jd/jd/202304/t20230407_1353388_ext.html", "必要性、可行性、风险可控性，以及需求、方案、运营、投融资和风险分析框架。"],
            ["2", "财政部会计司《企业内部控制应用指引第11号：工程项目》解读", "https://kjs.mof.gov.cn/zhengcejiedu/201007/t20100712_327373.htm", "项目建议书通常包括必要性依据、产品方案、建设规模地点、投资估算、资金筹措、进度安排和效益分析。"],
            ["3", "《党政机关公文格式》GB/T 9704-2012 公开转载件", "https://yb.jsut.edu.cn/2024/0610/c7386a181105/page.htm", "正式中文文档常用 A4 页面、版心和标题正文排版参考。"],
            ["4", "总决赛定稿-不亦“乐”乎（图片版）.pptx", "/Users/shishangbo/Desktop/总决赛定稿-不亦“乐”乎（图片版）.pptx", "统一项目背景口径，采用政策课程背景、学生参与弱、音乐感知难、课堂反馈慢等表述。"],
            ["5", "项目本地资料：README.md、智能体说明.md、DEPLOYMENT.md、智能体升级目标推进计划.md", "/Users/shishangbo/codex/第一版", "提炼项目定位、核心功能、技术架构、部署方式、升级路线和当前能力边界。"],
        ],
        [800, 3000, 3600, 2900],
    )

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    build_doc()
