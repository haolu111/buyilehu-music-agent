from __future__ import annotations

from datetime import date
from pathlib import Path
import textwrap

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
DESKTOP = Path.home() / "Desktop"
RAW_DIR = ROOT / "tmp" / "technical-manual-assets" / "ops-raw"
LEGACY_RAW_DIR = ROOT / "tmp" / "technical-manual-assets" / "raw"
EDITED_DIR = ROOT / "tmp" / "technical-manual-assets" / "edited"
OUTPUT = DESKTOP / "不亦乐乎音乐课堂游戏生成智能体技术手册_最新版.docx"

FONT_EAST_ASIA = "Microsoft YaHei"
FONT_HEADING = "SimHei"
FONT_LATIN = "Calibri"
INK = "172126"
MUTED = "667A77"
TEAL = "37A99B"
TEAL_DARK = "1E6F64"
GOLD = "D8A23C"
BLUE = "287E9B"
BLUE_DARK = "175B70"
FILL_LIGHT = "F8FBFA"
FILL_TEAL = "E8F6F3"
FILL_GOLD = "FFF5DA"
FILL_BLUE = "EAF5F8"
BORDER = "C9DDD8"
TABLE_HEADER = "DDEFEA"
TABLE_ALT = "F7FBFA"


SCREENSHOTS = [
    {
        "source": "main-auth.png",
        "name": "01-login",
        "title": "账号登录界面",
        "subtitle": "支持登录、注册、忘记密码三类入口，作品和账号按用户隔离。",
        "callouts": ["登录进入智能体", "注册需邮箱验证码", "忘记密码可重置"],
    },
    {
        "source": "01-home.png",
        "name": "02-home",
        "title": "首页与生成入口",
        "subtitle": "首页提供教案生成和直接生成两个主入口，教师无需先学习复杂配置。",
        "callouts": ["教案生成", "直接生成", "我的作品"],
    },
    {
        "source": "02-lesson.png",
        "name": "03-lesson",
        "title": "教案生成工作区",
        "subtitle": "上传或粘贴教案，可补充乐谱、音频和说明，先确认方案再生成。",
        "callouts": ["粘贴教案正文", "上传教案/乐谱/音频", "识别重点并确认方案"],
    },
    {
        "source": "03-direct.png",
        "name": "04-direct",
        "title": "直接生成工作区",
        "subtitle": "选择工具类型，描述课堂目标，可用快速或严格模式生成可打开作品。",
        "callouts": ["选择工具类型", "描述课堂目标", "快速/严格模式"],
    },
    {
        "source": "05-inspiration.png",
        "name": "05-inspiration",
        "title": "灵感助手操作区",
        "subtitle": "想法还不完整时，可先和灵感助手对话，再把建议放入生成框。",
        "callouts": ["提出初步想法", "获得玩法建议", "放入生成框"],
    },
    {
        "source": "04-library.png",
        "name": "06-library",
        "title": "我的作品与生成状态",
        "subtitle": "作品列表、生成进度和账号管理集中在“我的”页面，便于复用和继续修改。",
        "callouts": ["实时生成状态", "作品预览卡片", "账号管理"],
    },
    {
        "source": "06-revision.png",
        "name": "07-revision",
        "title": "增量修改工作区",
        "subtitle": "选择作品后，通过修改助手描述调整要求，只修改当前作品版本。",
        "callouts": ["输入修改要求", "发送理解", "应用修改"],
    },
]


def _pil_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "/System/Library/Fonts/Supplemental/Microsoft YaHei.ttf",
        "/System/Library/Fonts/Supplemental/SimHei.ttf" if bold else "/System/Library/Fonts/Supplemental/Microsoft YaHei.ttf",
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
        "/System/Library/Fonts/STHeiti Medium.ttc" if bold else "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/Supplemental/Songti.ttc",
        "/System/Library/Fonts/Supplemental/Arial Unicode.ttf",
    ]
    for candidate in candidates:
        try:
            return ImageFont.truetype(candidate, size=size)
        except OSError:
            continue
    return ImageFont.load_default()


def _wrap_text(text: str, chars: int) -> str:
    return "\n".join(textwrap.wrap(text, width=chars, break_long_words=False, replace_whitespace=False))


def _rounded_rect(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], radius: int, fill: str, outline: str | None = None, width: int = 1) -> None:
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def _draw_label(draw: ImageDraw.ImageDraw, xy: tuple[int, int], label: str, fill: str, text_fill: str = "#0F2F2B") -> None:
    font = _pil_font(28, bold=True)
    x, y = xy
    bbox = draw.textbbox((0, 0), label, font=font)
    w = bbox[2] - bbox[0] + 42
    h = bbox[3] - bbox[1] + 24
    _rounded_rect(draw, (x, y, x + w, y + h), 24, fill=fill, outline="#FFFFFF", width=2)
    draw.text((x + 21, y + 10), label, font=font, fill=text_fill)


def build_edited_screenshot(item: dict[str, object]) -> Path:
    source = RAW_DIR / str(item["source"])
    if not source.exists():
        source = LEGACY_RAW_DIR / str(item["source"])
    img = Image.open(source).convert("RGB")
    canvas_w, canvas_h = 1600, 1030
    canvas = Image.new("RGB", (canvas_w, canvas_h), "#F7FBFA")
    draw = ImageDraw.Draw(canvas)

    title_font = _pil_font(50, bold=True)
    sub_font = _pil_font(28)
    badge_font = _pil_font(24, bold=True)

    draw.rounded_rectangle((-28, -32, canvas_w + 28, 174), radius=34, fill="#E6F5F1")
    draw.rectangle((64, 40, 76, 126), fill="#37A99B")
    draw.text((96, 40), str(item["title"]), font=title_font, fill="#122B2A")
    draw.text((96, 106), _wrap_text(str(item["subtitle"]), 42), font=sub_font, fill="#4B6260")
    draw.rounded_rectangle((1290, 46, 1534, 112), radius=28, fill="#FFFFFF", outline="#CFE4DF", width=2)
    draw.text((1332, 65), "操作截图", font=badge_font, fill="#2B7F73")

    # Drop shadow and screenshot frame.
    frame_x, frame_y, frame_w = 64, 196, 1472
    scale = frame_w / img.width
    frame_h = int(img.height * scale)
    img = img.resize((frame_w, frame_h), Image.Resampling.LANCZOS)
    shadow = Image.new("RGBA", (frame_w + 22, frame_h + 22), (0, 0, 0, 0))
    shadow_draw = ImageDraw.Draw(shadow)
    shadow_draw.rounded_rectangle((12, 12, frame_w + 12, frame_h + 12), radius=32, fill=(24, 78, 72, 40))
    canvas.paste(shadow.convert("RGB"), (frame_x + 8, frame_y + 8))
    mask = Image.new("L", (frame_w, frame_h), 0)
    mask_draw = ImageDraw.Draw(mask)
    mask_draw.rounded_rectangle((0, 0, frame_w, frame_h), radius=28, fill=255)
    canvas.paste(img, (frame_x, frame_y), mask)
    draw.rounded_rectangle((frame_x, frame_y, frame_x + frame_w, frame_y + frame_h), radius=28, outline="#C8DAD6", width=3)

    y = 922
    for index, callout in enumerate(item["callouts"], start=1):
        x = 72 + (index - 1) * 495
        _draw_label(draw, (x, y), f"{index}. {callout}", "#FFF4D9" if index == 2 else "#E7F5F1")

    EDITED_DIR.mkdir(parents=True, exist_ok=True)
    out = EDITED_DIR / f"{item['name']}.png"
    canvas.save(out)
    return out


def build_usage_flow() -> Path:
    EDITED_DIR.mkdir(parents=True, exist_ok=True)
    out = EDITED_DIR / "usage-flow.png"
    canvas = Image.new("RGB", (1600, 720), "#F8FCFB")
    draw = ImageDraw.Draw(canvas)
    title_font = _pil_font(46, bold=True)
    step_font = _pil_font(30, bold=True)
    body_font = _pil_font(22)
    draw.text((70, 46), "智能体操作主流程", font=title_font, fill="#122B2A")
    steps = [
        ("1 登录", "输入账号进入平台"),
        ("2 选择入口", "教案生成或直接生成"),
        ("3 灵感整理", "想法不完整时先对话"),
        ("4 填写材料", "教案、乐谱、音频或目标"),
        ("5 确认方案", "检查重点、玩法和模式"),
        ("6 生成作品", "等待任务进度完成"),
        ("7 打开作品", "预览、试玩或下载"),
        ("8 增量修改", "只调整当前作品版本"),
    ]
    x0, y0, box_w, box_h, gap_x, gap_y = 70, 155, 330, 150, 46, 46
    for i, (title, body) in enumerate(steps):
        row, col = divmod(i, 4)
        x = x0 + col * (box_w + gap_x)
        y = y0 + row * (box_h + gap_y)
        fill = "#E7F5F1" if i % 2 == 0 else "#FFF6DF"
        outline = "#A9D6CC" if i % 2 == 0 else "#E7C879"
        _rounded_rect(draw, (x, y, x + box_w, y + box_h), 24, fill=fill, outline=outline, width=3)
        draw.text((x + 24, y + 26), title, font=step_font, fill="#173B36")
        draw.text((x + 24, y + 82), _wrap_text(body, 13), font=body_font, fill="#516460")
        if i < len(steps) - 1 and col < 3:
            ax = x + box_w + 9
            ay = y + box_h // 2
            draw.line((ax, ay, ax + gap_x - 18, ay), fill="#68AFA3", width=5)
            draw.polygon([(ax + gap_x - 18, ay - 11), (ax + gap_x + 2, ay), (ax + gap_x - 18, ay + 11)], fill="#68AFA3")
    note_font = _pil_font(26, bold=True)
    draw.rounded_rectangle((70, 560, 1530, 660), radius=24, fill="#FFFFFF", outline="#D6E2DF", width=2)
    draw.text((110, 588), "原则：先把课堂想法整理清楚，再生成作品；后续修改默认只调整当前作品版本。", font=note_font, fill="#1F6F64")
    draw.text((110, 628), "灵感助手负责整理想法，修改助手负责理解修改要求，真正生成和保存由后台任务完成。", font=_pil_font(22), fill="#4B6260")
    canvas.save(out)
    return out


def build_architecture_flow() -> Path:
    EDITED_DIR.mkdir(parents=True, exist_ok=True)
    out = EDITED_DIR / "architecture-flow.png"
    canvas = Image.new("RGB", (1600, 860), "#F8FCFB")
    draw = ImageDraw.Draw(canvas)
    title_font = _pil_font(46, bold=True)
    box_font = _pil_font(28, bold=True)
    body_font = _pil_font(21)
    draw.text((70, 44), "最新版技术架构", font=title_font, fill="#122B2A")
    boxes = [
        ("前端操作台", "登录、首页、教案生成、直接生成、灵感助手、我的作品"),
        ("FastAPI 后端", "账号、任务、作品、上传、生成接口"),
        ("教案设计大脑", "ChatECNU 优先，DeepSeek 备用，支持灵感整理"),
        ("生成编排层", "解析需求、形成方案、创建任务、记录进度"),
        ("执行与渲染层", "模板实例化或 OpenCode 任务包，生成独立网页"),
        ("修改与版本层", "修改助手、增量 patch、版本快照、当前编辑版本"),
        ("课堂作品", "可预览、下载、继续修改、复用"),
    ]
    positions = [
        (70, 150), (420, 150), (770, 150), (1120, 150),
        (245, 480), (595, 480), (945, 480),
    ]
    box_w, box_h = 300, 190
    for i, ((title, body), (x, y)) in enumerate(zip(boxes, positions)):
        fill = "#E7F5F1" if i in {0, 3, 6} else "#FFFFFF"
        _rounded_rect(draw, (x, y, x + box_w, y + box_h), 26, fill=fill, outline="#CFE0DC", width=3)
        draw.text((x + 24, y + 32), title, font=box_font, fill="#173B36")
        draw.text((x + 24, y + 88), _wrap_text(body, 13), font=body_font, fill="#526662")
    arrow_color = "#68AFA3"
    arrows = [
        ((370, 245), (420, 245)),
        ((720, 245), (770, 245)),
        ((1070, 245), (1120, 245)),
        ((1270, 340), (395, 480)),
        ((545, 575), (595, 575)),
        ((895, 575), (945, 575)),
    ]
    for (x1, y1), (x2, y2) in arrows:
        draw.line((x1, y1, x2 - 18, y2), fill=arrow_color, width=6)
        draw.polygon([(x2 - 18, y2 - 12), (x2 + 2, y2), (x2 - 18, y2 + 12)], fill=arrow_color)
    draw.rounded_rectangle((70, 720, 1530, 806), radius=22, fill="#FFF6DF", outline="#E8CE8B", width=2)
    draw.text((108, 746), "用户只面对浏览器操作；识谱、音频处理、代码生成、存储和验收统一在服务端或本机后端完成。", font=_pil_font(25, bold=True), fill="#6E4C00")
    canvas.save(out)
    return out


def set_run_font(run, size: float, *, bold: bool = False, color: str = INK) -> None:
    run.font.name = FONT_LATIN
    run._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING if bold else FONT_EAST_ASIA)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def set_paragraph(paragraph, *, before: float = 0, after: float = 6, line: float = 1.25, align: WD_ALIGN_PARAGRAPH | None = None) -> None:
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def set_paragraph_bottom_border(paragraph, color: str = TEAL, size: str = "12", space: str = "4") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), space)
    bottom.set(qn("w:color"), color)


def add_spacer(doc: Document, height: float = 5) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=0, after=height, line=1)


def shade_cell(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "8") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_border_edges(cell, *, top: str | None = None, bottom: str | None = None, left: str | None = None, right: str | None = None, size: str = "12") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge, color in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        if color is None:
            element.set(qn("w:val"), "nil")
        else:
            element.set(qn("w:val"), "single")
            element.set(qn("w:sz"), size)
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), color)


def set_cell_margins(cell, top: int = 100, start: int = 140, bottom: int = 100, end: int = 140) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    margins = tc_pr.find(qn("w:tcMar"))
    if margins is None:
        margins = OxmlElement("w:tcMar")
        tc_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        element = margins.find(qn(f"w:{name}"))
        if element is None:
            element = OxmlElement(f"w:{name}")
            margins.append(element)
        element.set(qn("w:w"), str(value))
        element.set(qn("w:type"), "dxa")


def set_column_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_width(table, width_dxa: int = 9360) -> None:
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "120")
    tbl_ind.set(qn("w:type"), "dxa")


def add_page_number(paragraph) -> None:
    field_run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    field_run._r.append(fld_char)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    field_run._r.append(instr)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "end")
    field_run._r.append(fld_char)
    run = paragraph.add_run(" / ")
    set_run_font(run, 9, color=MUTED)
    field_run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    field_run._r.append(fld_char)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "NUMPAGES"
    field_run._r.append(instr)
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "end")
    field_run._r.append(fld_char)


def configure_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.92)
    section.bottom_margin = Inches(0.82)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.35)
    section.different_first_page_header_footer = True

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT_LATIN
    normal._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_EAST_ASIA)
    normal.font.size = Pt(11)

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = styles[name]
        style.font.name = FONT_LATIN
        style._element.rPr.rFonts.set(qn("w:ascii"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:hAnsi"), FONT_LATIN)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), FONT_HEADING)
    styles["Heading 1"].font.size = Pt(16)
    styles["Heading 1"].font.color.rgb = RGBColor.from_string("000000")
    styles["Heading 2"].font.size = Pt(13)
    styles["Heading 2"].font.color.rgb = RGBColor.from_string("000000")
    styles["Heading 3"].font.size = Pt(12)
    styles["Heading 3"].font.color.rgb = RGBColor.from_string("000000")

    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(header, after=0, line=1)
    set_run_font(header.add_run("不亦乐乎音乐课堂游戏生成智能体技术手册"), 9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_page_number(footer)


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph(style=f"Heading {level}")
    before = 8 if level == 1 else 8 if level == 2 else 6
    after = 8 if level == 1 else 5 if level == 2 else 3
    set_paragraph(p, before=before, after=after)
    run = p.add_run(text)
    color = "000000"
    size = 16 if level == 1 else 13 if level == 2 else 12
    set_run_font(run, size, bold=True, color=color)


def add_body(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, after=7, line=1.28)
    set_run_font(p.add_run(text), 11)


def add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.32)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    set_paragraph(p, after=4, line=1.22)
    set_run_font(p.add_run(text), 10.5)


def add_number(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.34)
    p.paragraph_format.first_line_indent = Inches(-0.16)
    set_paragraph(p, after=4, line=1.22)
    set_run_font(p.add_run(text), 10.5)


def add_callout(doc: Document, text: str, fill: str = FILL_TEAL) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.18)
    set_paragraph(p, before=4, after=6, line=1.25)
    set_run_font(p.add_run(text), 11, bold=False, color=INK)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table)
    table.rows[0].height = None
    tr_pr = table.rows[0]._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_column_width(cell, widths[i])
        shade_cell(cell, TABLE_HEADER)
        set_cell_border(cell)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        set_paragraph(p, after=0)
        set_run_font(p.add_run(header), 10.5, bold=True, color=BLUE_DARK)
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cell = cells[i]
            set_column_width(cell, widths[i])
            if row_index % 2 == 1:
                shade_cell(cell, TABLE_ALT)
            set_cell_border(cell)
            set_cell_margins(cell, top=120, start=150, bottom=120, end=150)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            set_paragraph(p, after=0, line=1.22)
            set_run_font(p.add_run(value), 10)
    add_spacer(doc, 7)


def add_table_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, before=14, after=16, line=1.1)
    set_run_font(p.add_run(text), 10.5, color=BLUE_DARK)


def add_minimal_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[int]) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        set_column_width(cell, widths[i])
        shade_cell(cell, "EAF2F8")
        set_cell_margins(cell, top=180, start=160, bottom=180, end=160)
        set_cell_border_edges(cell, top=BLUE_DARK, bottom=BLUE_DARK, left=None, right=None, size="14")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph(p, after=0, line=1.2)
        set_run_font(p.add_run(header), 10.5, bold=True, color="000000")
    for row_index, row in enumerate(rows):
        cells = table.add_row().cells
        is_last = row_index == len(rows) - 1
        for i, value in enumerate(row):
            cell = cells[i]
            set_column_width(cell, widths[i])
            set_cell_margins(cell, top=165, start=160, bottom=165, end=160)
            set_cell_border_edges(cell, top=None, bottom=BLUE_DARK if is_last else None, left=None, right=None, size="14")
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 2 else WD_ALIGN_PARAGRAPH.LEFT
            set_paragraph(p, after=0, line=1.45)
            set_run_font(p.add_run(value), 10.2, color="000000")
    add_spacer(doc, 13)


def add_image_figure(doc: Document, image_path: Path, caption: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, before=5, after=3)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.45))
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(cap, before=0, after=9, line=1.1)
    set_run_font(cap.add_run(caption), 9.5, color=MUTED)


def cover(doc: Document) -> None:
    p = doc.add_paragraph()
    set_paragraph(p, before=78, after=12, align=WD_ALIGN_PARAGRAPH.CENTER)
    logo = ROOT / "app" / "static" / "assets" / "buyilehu-logo-transparent.png"
    if logo.exists():
        p.add_run().add_picture(str(logo), width=Inches(2.0))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, before=28, after=4, line=1.08)
    set_run_font(p.add_run("不亦乐乎"), 34, bold=True, color="000000")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, after=8, line=1.08)
    set_run_font(p.add_run("音乐课堂游戏生成智能体技术手册"), 26, bold=True, color="000000")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph(p, before=10, after=52)
    set_run_font(p.add_run("面向中小学音乐课堂的智能体操作平台"), 14, color="000000")
    for line in [
        f"版本：最新版（{date.today().isoformat()}）",
        "适用对象：音乐教师、教研员、演示评审人员",
        "本地访问：http://127.0.0.1:8000",
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_paragraph(p, after=6)
        set_run_font(p.add_run(line), 11.5, color="000000")
    doc.add_page_break()


def toc(doc: Document) -> None:
    add_heading(doc, "目录", 1)
    entries = [
        ("目录", "2"),
        ("智能体概述", "3"),
        ("功能模块", "4"),
        ("使用流程", "6"),
        ("技术架构", "9"),
        ("运行与部署", "10"),
        ("资源规模", "11"),
        ("界面展示", "12"),
        ("操作注意事项", "19"),
    ]
    for title, page in entries:
        p = doc.add_paragraph()
        p.paragraph_format.tab_stops.add_tab_stop(Inches(6.25), WD_TAB_ALIGNMENT.RIGHT, WD_TAB_LEADER.DOTS)
        set_paragraph(p, after=5, line=1.2)
        set_run_font(p.add_run(title), 11.5, color="000000")
        set_run_font(p.add_run("\t" + page), 11.5, color="000000")
    doc.add_page_break()


def overview(doc: Document) -> None:
    add_heading(doc, "智能体概述", 1)
    add_body(doc, "不亦乐乎音乐课堂游戏生成智能体不是单纯的 AI 聊天页，而是一套面向中小学音乐课堂的生成工作台。教师可以上传教案、乐谱和音频，也可以直接描述课堂目标，系统会把需求转成可打开、可预览、可继续修改的课堂作品。")
    add_body(doc, "智能体主界面只负责登录、需求输入、方案确认、生成进度和作品管理；学生或课堂展示使用的是生成后的独立网页作品。")
    add_callout(doc, "核心价值：把教师的音乐课堂想法，转化为能听、能唱、能玩、能复用的互动课堂作品。")
    add_heading(doc, "适用场景", 2)
    for text in [
        "课前备课：将教案中的教学重点、课堂环节和音乐材料转成可操作活动。",
        "课中互动：生成可播放、可点击、可拖拽或可闯关的课堂网页。",
        "课后复用：在“我的作品”中打开历史作品，继续修改或再次使用。",
    ]:
        add_bullet(doc, text)


def modules(doc: Document) -> None:
    add_heading(doc, "功能模块", 1)
    sections = [
        ("账号登录", "登录模块提供登录、注册、忘记密码三个入口。用户作品与账号隔离，新用户通过邮箱验证码完成注册，忘记密码时可通过邮箱验证码重置。"),
        ("首页", "首页是教师进入智能体后的第一操作页，集中提供“教案生成”“直接生成”和“我的作品”等入口，适合演示和日常备课。"),
        ("教案生成", "教案生成用于把完整课例转化为课堂作品。教师可以粘贴或上传教案，也可以补充乐谱、音频和具体要求。系统会先识别教学重点和方案，再进入生成环节。"),
        ("直接生成", "直接生成适合从一句需求快速开始。教师选择工具类型，填写曲目、年级、音乐要素和课堂目标后，可以使用快速模式或严格模式生成作品；想法还不完整时，可先使用页面左侧的灵感助手整理玩法方向，再放入生成框。"),
        ("我的作品", "我的作品用于查看生成状态、打开预览、继续修改和管理账号。作品按用户保存，便于后续复用；选择作品后可进入修改助手，用自然语言说明调整要求并应用到当前编辑版本。"),
        ("生成任务", "生成任务由后台队列执行，前端显示实时进度。任务完成后，结果以可打开的独立网页作品形式进入作品列表。"),
    ]
    for title, body in sections:
        add_heading(doc, title, 2)
        add_body(doc, body)


def usage(doc: Document, usage_flow: Path) -> None:
    add_heading(doc, "使用流程", 1)
    add_image_figure(doc, usage_flow, "图 1：教师使用智能体的主流程")
    add_heading(doc, "教案生成流程", 2)
    for text in [
        "登录后进入首页，选择“教案生成”。",
        "粘贴教案正文，或上传 Word、TXT、Markdown、PDF 格式教案。",
        "按需要补充乐谱图片/PDF、MIDI、MusicXML、文字谱或音频。",
        "点击识别与分析后，查看系统提炼出的音乐重点、投放环节和玩法建议。",
        "确认方案后生成作品，完成后到“我的作品”中预览、下载或继续修改。",
    ]:
        add_number(doc, text)
    add_heading(doc, "直接生成流程", 2)
    for text in [
        "进入“直接生成”，选择聆听工具、表现闯关、创造拼图或音乐小游戏。",
        "如果想法还不完整，可先在左侧“灵感助手”输入曲目、年级和教学目标，让系统整理玩法建议。",
        "点击“放入生成框”后，继续补充曲目、音乐要素、学生任务和评价要求。",
        "写清曲目、年级、音乐要素、学生任务和通关或评价要求。",
        "选择快速模式或严格模式：快速模式适合先出第一版，严格模式适合定稿前验收。",
        "提交后等待生成状态更新，完成后在“我的作品”中打开结果。",
    ]:
        add_number(doc, text)
    add_heading(doc, "我的作品与继续修改流程", 2)
    for text in [
        "进入“我的作品”，选择需要调整的作品，点击“继续修改”。",
        "在作品详情页查看“当前编辑版本”，确认后续修改会应用到当前版本。",
        "在“修改助手”中用自然语言说明要改的内容，例如降低第一关难度、放大按钮、调整反馈提示或保留某段音乐材料。",
        "先点击“发送理解”，查看修改助手整理出的修改说明。",
        "确认无误后点击“应用修改”，系统会保存新版本；旧版本只用于打开查看，不作为后续修改对象。",
    ]:
        add_number(doc, text)


def architecture(doc: Document, arch_flow: Path) -> None:
    add_heading(doc, "技术架构", 1)
    add_image_figure(doc, arch_flow, "图 2：最新版智能体技术架构")
    sections = [
        ("前端", "主站采用 HTML、CSS 与 JavaScript 实现，模板配置台使用 React 与 Vite。界面负责登录、材料上传、需求输入、任务状态展示和作品入口。"),
        ("后端", "后端采用 FastAPI，提供登录、上传、作品、任务、生成、修改和部署状态等接口，负责把用户操作转化为可执行任务。"),
        ("模型层", "ChatECNU 作为教案设计大脑和灵感整理模型，DeepSeek 可作为备用模型。模型不可用时，系统使用本地规则兜底，保证基本流程可继续演示。"),
        ("任务层", "JobQueue 负责排队、状态、超时、重试、事件日志和恢复，前端通过任务状态向用户展示生成进度。"),
        ("音乐处理", "系统可接入 Basic Pitch、MIDI 处理、SoundFont、本地音频渲染以及可选 OMR/OCR，用于处理乐谱、音频和音乐材料。"),
        ("生成执行", "成熟需求可直接实例化为作品；复杂需求可准备 OpenCode 任务包或调用 OpenCode，生成独立课堂网页。"),
        ("增量修改", "修改助手会先把教师自然语言要求整理成修改说明，再通过当前作品规格、文件路径和版本上下文执行增量修改，避免每次都从零重做。"),
        ("质量保障", "执行编排层检查音乐逻辑、教案贴合、页面可用性、结构配置和版本快照，确保生成和修改结果可以用于课堂展示。"),
    ]
    for title, body in sections:
        add_heading(doc, title, 2)
        add_body(doc, body)


def deployment(doc: Document) -> None:
    add_heading(doc, "运行与部署", 1)
    add_body(doc, "本项目既可以在教师或开发者电脑本地运行，也可以部署为服务器网址。正式演示或比赛提交时，推荐把识谱、音频处理、代码生成和文件存储放在服务器端，用户只通过浏览器访问。")
    add_heading(doc, "本地启动", 2)
    for text in [
        "后端：创建 Python 虚拟环境，安装 requirements.txt，运行 uvicorn app.main:app --reload。",
        "前端：frontend 目录中运行 npm install 与 npm run dev，开发地址为 http://127.0.0.1:5173/template-console/。",
        "主站：浏览器打开 http://127.0.0.1:8000；本机可使用 /dev 进入开发会话。",
    ]:
        add_bullet(doc, text)
    add_heading(doc, "服务器部署", 2)
    for text in [
        "使用 Dockerfile.server 与 docker-compose.yml 构建服务，端口由 MUSIC_AGENT_PORT 控制。",
        "生产环境需配置 SMTP，用于注册和忘记密码验证码。",
        "如需乐谱识别，服务器配置 Audiveris、简谱 OMR 或 Tesseract OCR。",
        "部署完成后通过 /api/deployment/status 查看比赛或公开演示所需的 readiness 状态。",
    ]:
        add_bullet(doc, text)


def _count_files(path: Path, pattern: str = "*") -> int:
    return sum(1 for item in path.rglob(pattern) if item.is_file()) if path.exists() else 0


def _count_dirs(path: Path) -> int:
    return sum(1 for item in path.iterdir() if item.is_dir()) if path.exists() else 0


def resource_scale(doc: Document) -> None:
    service_count = _count_files(ROOT / "app" / "services", "*.py")
    static_count = _count_files(ROOT / "app" / "static" / "assets")
    frontend_count = _count_files(ROOT / "frontend" / "src")
    test_count = _count_files(ROOT / "tests", "test_*.py")
    activity_pack_count = _count_dirs(ROOT / "app" / "static" / "assets" / "game-packs")
    mascot_count = _count_files(ROOT / "app" / "static" / "assets" / "cute-mascots")

    add_heading(doc, "资源规模", 1)
    add_body(doc, "根据当前目录内容，平台已经具备较完整的音乐课堂智能体资源基础：")

    add_table_caption(doc, "表 1：平台资源规模")
    add_minimal_table(
        doc,
        ["资源类型", "数量", "说明"],
        [
            ("服务模块", f"{service_count} 个", "账号、上传、任务、生成、修改、音乐处理、部署状态等后端能力"),
            ("静态资源", f"{static_count} 个", f"界面素材、互动活动包、音频与视觉资源，其中包含 {activity_pack_count} 组本地互动活动包和 {mascot_count} 个吉祥物素材"),
            ("前端源码", f"{frontend_count} 个", "主站页面、模板配置台、生成入口、作品管理与交互状态管理"),
            ("测试文件", f"{test_count} 个", "接口、任务队列、生成流程、音乐处理与版本修改等关键路径验证"),
        ],
        [2200, 2100, 5060],
    )

    p = doc.add_paragraph()
    set_paragraph(p, before=8, after=8)
    set_run_font(p.add_run("示例资源："), 11, color="000000")

    add_table_caption(doc, "表 2：示例资源列表")
    add_minimal_table(
        doc,
        ["类型", "示例"],
        [
            ("操作入口", "登录、首页、教案生成、直接生成、灵感助手、我的作品、继续修改、账号管理"),
            ("输入材料", "Word 教案、PDF 教案、乐谱图片、MusicXML、MIDI、MP3/WAV 音频、手写补充说明"),
            ("生成能力", "教案生成、直接生成、快速模式、严格模式、灵感整理、作品预览、下载复用、增量修改"),
            ("工程能力", "FastAPI 后端、JobQueue 任务队列、ChatECNU/DeepSeek 模型接入、本地规则兜底、版本快照"),
        ],
        [2200, 7160],
    )


def interface_showcase(doc: Document, figures: list[Path]) -> None:
    add_heading(doc, "界面展示", 1)
    captions = [
        ("登录界面", ["打开平台网址。", "输入账号密码登录；新用户选择注册并获取邮箱验证码。", "忘记密码时通过邮箱验证码重置。"]),
        ("首页界面", ["登录后进入首页。", "选择“教案生成”或“直接生成”。", "也可以进入“我的”查看作品。"]),
        ("教案生成界面", ["粘贴教案正文或上传教案文件。", "按需上传乐谱、音频或补充要求。", "先确认智能体识别出的重点和方案，再生成作品。"]),
        ("直接生成界面", ["选择工具类型。", "描述曲目、年级、音乐要素和学生任务。", "选择快速或严格模式后点击生成。"]),
        ("直接生成中的灵感助手", ["想法还不完整时，在直接生成页面左侧输入曲目、年级和初步目标。", "查看灵感助手整理的玩法方向、课堂任务和可生成描述。", "点击“放入生成框”，再在右侧生成区补充规则并生成作品。"]),
        ("我的作品界面", ["查看当前生成状态。", "在作品卡片中打开预览。", "进入账号页可管理密码或退出登录。"]),
        ("我的作品中的增量修改", ["在“我的作品”中选择作品并进入继续修改。", "在修改助手中输入要调整的内容，并点击“发送理解”。", "确认修改说明后点击“应用修改”，修改会保存到当前编辑版本。"]),
    ]
    for index, (figure, (title, steps)) in enumerate(zip(figures, captions), start=1):
        add_heading(doc, title, 2)
        add_image_figure(doc, figure, f"图 {index + 2}：{title}")
        p = doc.add_paragraph()
        set_paragraph(p, before=0, after=3)
        set_run_font(p.add_run("操作步骤："), 10.5, bold=True, color="000000")
        for step in steps:
            add_bullet(doc, step)
        if index < len(figures):
            doc.add_page_break()


def notes(doc: Document) -> None:
    add_heading(doc, "操作注意事项", 1)
    for text in [
        "教案生成适合正式课例：先让智能体理解整节课，再确认玩法，结果更贴近教学目标。",
        "直接生成适合快速试做：一句需求即可生成第一版，后续可在作品中继续修改。",
        "灵感助手只负责整理想法和玩法建议，不直接生成网页；整理好的建议需要放入生成框后再提交。",
        "增量修改默认作用在当前编辑版本上，旧版本只用于打开查看，避免把后续修改误应用到历史版本。",
        "上传材料越具体，生成结果越稳定；曲目、学段、音乐要素、学生动作和评价规则建议写清楚。",
        "若用于公开演示，建议提前完成服务器 readiness 检查，确保账号验证码、输出目录、音频工具和生成执行环境可用。",
        "学生端只需要打开作品网页，不需要安装 Python、OpenCode、识谱工具或模型环境。",
    ]:
        add_bullet(doc, text)
    add_callout(doc, "本手册为最新版操作技术手册，依据 2026-05-28 当前工作区界面和代码生成。")


def build() -> Path:
    EDITED_DIR.mkdir(parents=True, exist_ok=True)
    figures = [build_edited_screenshot(item) for item in SCREENSHOTS]
    usage_flow = build_usage_flow()
    arch_flow = build_architecture_flow()

    doc = Document()
    configure_doc(doc)
    cover(doc)
    toc(doc)
    overview(doc)
    modules(doc)
    doc.add_page_break()
    usage(doc, usage_flow)
    doc.add_page_break()
    architecture(doc, arch_flow)
    doc.add_page_break()
    deployment(doc)
    doc.add_page_break()
    resource_scale(doc)
    doc.add_page_break()
    interface_showcase(doc, figures)
    doc.add_page_break()
    notes(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
